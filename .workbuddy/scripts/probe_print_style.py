# -*- coding: utf-8 -*-
"""探查学生版 docx 现状：颜色 / 字体 / 底纹 / 图片 / 页面设置 / 页眉页脚"""
import sys, io, zipfile, re
from collections import Counter, defaultdict
from pathlib import Path
import docx
from docx.shared import Pt, RGBColor
from PIL import Image

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

ROOT = Path(r"C:\Obsidion\妙妙屋\00-首页\题组Word\习题书\学生版")
NS = {
    'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
}


def hexc(v):
    if not v or v == 'auto':
        return 'auto'
    return v.upper()


def scan(path: Path):
    d = docx.Document(str(path))
    z = zipfile.ZipFile(str(path))
    xml = z.read('word/document.xml').decode('utf-8', 'replace')

    colors = Counter()
    highlights = Counter()
    shd = Counter()
    fonts = Counter()
    sizes = Counter()

    # run 级
    for p in d.paragraphs:
        for r in p.runs:
            try:
                c = r.font.color
                colors[hexc(c.rgb if c and c.type is not None and c.rgb else 'auto')] += 1
            except Exception:
                colors['ERR'] += 1
            if r.font.highlight_color is not None:
                highlights[str(r.font.highlight_color)] += 1
            fonts[r.font.name or 'inherit'] += 1
            if r.font.size:
                sizes[round(r.font.size.pt, 1)] += 1
    for t in d.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for r in p.runs:
                        try:
                            c = r.font.color
                            colors[hexc(c.rgb if c and c.type is not None and c.rgb else 'auto')] += 1
                        except Exception:
                            pass
                        if r.font.highlight_color is not None:
                            highlights[str(r.font.highlight_color)] += 1

    # XML 级：底纹 & 段落底纹 & 边框颜色
    for m in re.finditer(r'<w:shd[^>]*w:fill="([0-9A-Fa-f]{6})"', xml):
        shd[m.group(1).upper()] += 1
    border_colors = Counter(re.findall(r'<w:(?:top|left|bottom|right|insideH|insideV)[^>]*w:color="([0-9A-Fa-f]{6})"', xml))
    tbl_borders = re.findall(r'<w:tblBorders>.*?</w:tblBorders>', xml, re.S)
    tbl_style = len(re.findall(r'<w:tbl>', xml))

    # 页面设置
    sec = d.sections[0]
    page = {
        'w_cm': round(sec.page_width.cm, 2) if sec.page_width else None,
        'h_cm': round(sec.page_height.cm, 2) if sec.page_height else None,
        'mt_cm': round(sec.top_margin.cm, 2),
        'mb_cm': round(sec.bottom_margin.cm, 2),
        'ml_cm': round(sec.left_margin.cm, 2),
        'mr_cm': round(sec.right_margin.cm, 2),
    }

    # 页眉页脚
    has_header = any(not s.header.is_linked_to_previous and s.header.paragraphs for s in d.sections)
    has_footer = any(not s.footer.is_linked_to_previous and s.footer.paragraphs for s in d.sections)
    footer_txt = ''
    if d.sections[0].footer.paragraphs:
        footer_txt = '|'.join(p.text.strip() for p in d.sections[0].footer.paragraphs)[:60]

    # 图片分析
    media = [n for n in z.namelist() if n.startswith('word/media/')]
    img_color = 0
    img_gray = 0
    img_dark = 0
    img_trans = 0
    for mn in media:
        try:
            im = Image.open(io.BytesIO(z.read(mn)))
            if im.mode in ('P',):
                im = im.convert('RGBA')
            if im.mode in ('LA', 'RGBA'):
                # 判断是否透明 + 是否彩色
                a = im.getchannel('A')
                if a.getextrema()[0] < 250:
                    img_trans += 1
                rgb = im.convert('RGB')
            else:
                rgb = im.convert('RGB')
            # 采样判断彩色度
            small = rgb.resize((min(120, im.width), min(120, im.height)))
            px = list(small.getdata())
            sat = max(abs(r - g) + abs(g - b) + abs(r - b) for r, g, b in px[:4000])
            if sat > 30:
                img_color += 1
            else:
                img_gray += 1
            # 暗背景检测（白底黑字图 vs 深色底图）
            lum = sum(0.299 * r + 0.587 * g + 0.114 * b for r, g, b in px) / len(px)
            if lum < 140:
                img_dark += 1
        except Exception as e:
            pass

    # 公式数
    omml = len(re.findall(r'<m:oMath[ >]', xml)) + len(re.findall(r'<m:oMathPara[ >]', xml))

    return dict(
        name=path.name,
        paras=len(d.paragraphs),
        tables=tbl_style,
        media=len(media),
        img_color=img_color, img_gray=img_gray, img_dark=img_dark, img_trans=img_trans,
        colors=dict(colors.most_common(8)),
        highlights=dict(highlights.most_common(5)),
        shd=dict(shd.most_common(5)),
        border_colors=dict(Counter(c.upper() for c in border_colors).most_common(5)),
        fonts=dict(fonts.most_common(5)),
        sizes=dict(sizes.most_common(6)),
        page=page,
        header=has_header, footer=has_footer, footer_txt=footer_txt,
        omml=omml,
        size_kb=round(path.stat().st_size / 1024),
    )


def main():
    files = sorted(ROOT.rglob('*.docx'))
    print(f"共 {len(files)} 个 docx\n")
    agg_color = Counter(); agg_hl = Counter(); agg_shd = Counter()
    agg_font = Counter(); agg_size = Counter(); agg_bc = Counter()
    tot = Counter()
    rows = []
    for f in files:
        try:
            r = scan(f)
        except Exception as e:
            print(f"[FAIL] {f.name}: {e}")
            continue
        rows.append(r)
        agg_color.update(r['colors']); agg_hl.update(r['highlights'])
        agg_shd.update(r['shd']); agg_font.update(r['fonts'])
        agg_size.update(r['sizes']); agg_bc.update(r['border_colors'])
        tot['paras'] += r['paras']; tot['tables'] += r['tables']
        tot['media'] += r['media']; tot['color_img'] += r['img_color']
        tot['gray_img'] += r['img_gray']; tot['dark_img'] += r['img_dark']
        tot['trans_img'] += r['img_trans']; tot['omml'] += r['omml']
        tot['header'] += 1 if r['header'] else 0
        tot['footer'] += 1 if r['footer'] else 0
        tot['size'] += r['size_kb']

    print("=" * 70)
    print("【总体】段落", tot['paras'], "| 表格", tot['tables'], "| 公式(OMML)", tot['omml'])
    print("图片", tot['media'], "= 彩色", tot['color_img'], "/ 灰度", tot['gray_img'],
          "| 其中暗背景", tot['dark_img'], "| 含透明", tot['trans_img'])
    print("有页眉", tot['header'], "| 有页脚", tot['footer'], "| 总体积", round(tot['size']/1024, 1), "MB")
    print("-" * 70)
    print("字体颜色 TOP:", agg_color.most_common(12))
    print("高亮 TOP:", agg_hl.most_common(8))
    print("底纹(shd) TOP:", agg_shd.most_common(10))
    print("边框色 TOP:", agg_bc.most_common(10))
    print("字体 TOP:", agg_font.most_common(10))
    print("字号 TOP:", agg_size.most_common(10))
    print("=" * 70)
    print(f"{'文件':<34}{'段':>5}{'表':>4}{'图':>4}{'彩图':>5}{'暗图':>5}{'公式':>5}{'页脚':>5}")
    for r in rows:
        rel = str(r['name'])
        print(f"{rel[:33]:<34}{r['paras']:>5}{r['tables']:>4}{r['media']:>4}"
              f"{r['img_color']:>5}{r['img_dark']:>5}{r['omml']:>5}{'Y' if r['footer'] else '-':>5}")
    print("\n页脚样例:", [r['footer_txt'] for r in rows[:3]])
    print("页面设置样例:", rows[0]['page'])
    # 收集所有不重复的页面设置
    pages = set(tuple(sorted(r['page'].items())) for r in rows)
    print("不同页面设置种类:", len(pages))
    for p in list(pages)[:5]:
        print("   ", dict(p))


if __name__ == '__main__':
    main()
