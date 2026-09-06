# -*- coding: utf-8 -*-
"""深挖：页脚内容 / styles.xml 默认 / ERR 颜色 / 边框位置 / 段落样式分布 / 行距"""
import sys, io, zipfile, re
from collections import Counter
from pathlib import Path
import docx

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
ROOT = Path(r"C:\Obsidion\妙妙屋\00-首页\题组Word\习题书\学生版")

f1 = ROOT / "第一篇-化学原理" / "1-热力学.docx"
f2 = ROOT / "来源索引.docx"
f3 = ROOT / "第三篇-有机化学" / "1-结构基础与波谱分析.docx"


def dump(name, p):
    print("=" * 72)
    print("###", name, p.name)
    z = zipfile.ZipFile(str(p))
    doc = docx.Document(str(p))

    # 1. 页脚 XML
    for n in z.namelist():
        if 'footer' in n and n.endswith('.xml'):
            x = z.read(n).decode('utf-8', 'replace')
            print(f"  [页脚 {n}] len={len(x)}")
            print("   ", re.sub(r'\s+', ' ', x)[:600])

    # 2. 页眉
    for n in z.namelist():
        if 'header' in n and n.endswith('.xml'):
            x = z.read(n).decode('utf-8', 'replace')
            print(f"  [页眉 {n}] {re.sub(r'<[^>]+>','',x)[:120]!r}")

    # 3. 边框 4A6577 上下文
    xml = z.read('word/document.xml').decode('utf-8', 'replace')
    for m in re.finditer(r'.{200}4A6577.{120}', xml):
        print("  [边框上下文]", re.sub(r'\s+', ' ', m.group(0))[-330:])
        break

    # 4. styles.xml 默认
    st = z.read('word/styles.xml').decode('utf-8', 'replace')
    m = re.search(r'<w:docDefaults>.*?</w:docDefaults>', st, re.S)
    if m:
        print("  [docDefaults]", re.sub(r'\s+', ' ', m.group(0))[:700])
    # Normal 样式
    m = re.search(r'<w:style [^>]*w:styleId="Normal".*?</w:style>', st, re.S)
    if m:
        print("  [Normal]", re.sub(r'\s+', ' ', m.group(0))[:500])
    # 样式 ID 列表
    sids = re.findall(r'<w:style [^>]*w:styleId="([^"]+)"[^>]*>\s*<w:name w:val="([^"]+)"', st)
    print("  [样式]", sids[:25])


def stylesurvey():
    print("\n" + "#" * 72)
    print("# 全库段落样式分布")
    agg = Counter()
    sz = Counter()
    spacing = Counter()
    for f in sorted(ROOT.rglob('*.docx')):
        d = docx.Document(str(f))
        for p in d.paragraphs:
            st = p.style.name if p.style else '?'
            agg[st] += 1
            pf = p.paragraph_format
            if pf.space_before is not None or pf.space_after is not None:
                spacing[(pf.space_before.pt if pf.space_before else None,
                         pf.space_after.pt if pf.space_after else None,
                         round(pf.line_spacing, 2) if isinstance(pf.line_spacing, float) else pf.line_spacing)] += 1
            for r in p.runs:
                if r.font.size:
                    sz[(st, round(r.font.size.pt, 1))] += 1
    print("段落样式 TOP:", agg.most_common(15))
    print("字号(按样式) TOP:", sz.most_common(15))
    print("段间距 TOP:", spacing.most_common(10))


def errcolor():
    print("\n" + "#" * 72)
    print("# ERR 颜色样本（无法读 rgb 的 run）")
    n = 0
    for f in sorted(ROOT.rglob('*.docx')):
        if n > 6:
            break
        z = zipfile.ZipFile(str(f))
        xml = z.read('word/document.xml').decode('utf-8', 'replace')
        for m in re.finditer(r'<w:color w:val="([^"]+)"', xml):
            v = m.group(1)
            if v != 'auto' and not re.match(r'^[0-9A-Fa-f]{6}$', v):
                print("  ", f.name, "->", repr(v))
                n += 1
                break
    # 统计所有 w:color 非 auto 值
    allc = Counter()
    for f in sorted(ROOT.rglob('*.docx')):
        z = zipfile.ZipFile(str(f))
        for part in ('word/document.xml', 'word/styles.xml'):
            try:
                xml = z.read(part).decode('utf-8', 'replace')
            except KeyError:
                continue
            for m in re.finditer(r'<w:color w:val="([^"]+)"', xml):
                allc[m.group(1)] += 1
    print("全库 w:color 取值:", allc.most_common(15))

    # 高亮
    hl = Counter()
    for f in sorted(ROOT.rglob('*.docx')):
        z = zipfile.ZipFile(str(f))
        xml = z.read('word/document.xml').decode('utf-8', 'replace')
        for m in re.finditer(r'<w:highlight w:val="([^"]+)"', xml):
            hl[m.group(1)] += 1
    print("全库 highlight:", hl.most_common(10))

    # 底纹
    sh = Counter()
    for f in sorted(ROOT.rglob('*.docx')):
        z = zipfile.ZipFile(str(f))
        for part in ('word/document.xml', 'word/styles.xml'):
            try:
                xml = z.read(part).decode('utf-8', 'replace')
            except KeyError:
                continue
            for m in re.finditer(r'<w:shd[^>]*w:fill="([^"]+)"', xml):
                sh[m.group(1)] += 1
    print("全库 shd fill:", sh.most_common(10))


if __name__ == '__main__':
    dump("样例A", f1)
    dump("样例B-索引", f2)
    dump("样例C-图多", f3)
    stylesurvey()
    errcolor()
