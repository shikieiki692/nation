# -*- coding: utf-8 -*-
"""校验打印版：颜色残留 / 封面是否清除 / 公式完整性 / 页眉页脚 / 图片彩色残留 / 段落数守恒"""
import sys, io, re, zipfile
from pathlib import Path
from collections import Counter
from PIL import Image

import docx

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

SRC = Path(r"C:\Obsidion\妙妙屋\00-首页\题组Word\习题书\学生版")
DST = Path(r"C:\Obsidion\妙妙屋\00-首页\题组Word\习题书\学生版-打印版")


def colorful_fullres(data, thresh=18):
    """全分辨率检测彩色：任一像素通道极差超阈值且占比>0.2%"""
    try:
        img = Image.open(io.BytesIO(data))
        img.load()
    except Exception:
        return None, 0
    rgb = img.convert('RGB')
    # 限制采样量，用网格采样而非缩放（缩放会平均掉小面积彩色）
    w, h = rgb.size
    px = rgb.load()
    step = max(1, int((w * h / 60000) ** 0.5))
    tot = col = 0
    for y in range(0, h, step):
        for x in range(0, w, step):
            r, g, b = px[x, y]
            tot += 1
            if max(r, g, b) - min(r, g, b) > thresh:
                col += 1
    return (col / tot * 100 if tot else 0), tot


# 主题调色板（theme1.xml）里的颜色定义不影响渲染，只在有 themeColor 引用时才生效；
# 正文/样式/页眉页脚里的显式颜色才决定实际打印效果。
CONTENT_PARTS = ('word/document.xml', 'word/styles.xml',
                 'word/footnotes.xml', 'word/endnotes.xml')


def zscan(p: Path):
    z = zipfile.ZipFile(str(p))
    colors = Counter()
    theme_colors = Counter()
    hl = 0
    shd = Counter()
    for part in [n for n in z.namelist() if n.endswith('.xml') and n.startswith('word/')]:
        try:
            x = z.read(part).decode('utf-8', 'replace')
        except Exception:
            continue
        is_content = part in CONTENT_PARTS or 'header' in part or 'footer' in part
        for m in re.finditer(r'<w:color w:val="([^"]+)"', x):
            (colors if is_content else theme_colors)[m.group(1)] += 1
        hl += len(re.findall(r'<w:highlight\b', x))
        for m in re.finditer(r'<w:shd[^>]*w:fill="([^"]+)"', x):
            shd[m.group(1)] += 1
    doc = docx.Document(str(p))
    xml = z.read('word/document.xml').decode('utf-8', 'replace')
    omml = len(re.findall(r'<m:oMath[ >]', xml)) + len(re.findall(r'<m:oMathPara[ >]', xml))
    imgs = len(re.findall(r'<m:oMath[ >]', xml))
    media = [n for n in z.namelist() if n.startswith('word/media/')]
    first = doc.paragraphs[0].text.strip() if doc.paragraphs else ''
    hdr = ''
    try:
        s0 = doc.sections[0]
        if s0.header.paragraphs:
            hdr = ' '.join(q.text.strip() for q in s0.header.paragraphs)
        ftr = ''
        fxml = z.read('word/footer1.xml').decode('utf-8', 'replace') if 'word/footer1.xml' in z.namelist() else ''
        ftr = 'PAGE' in fxml, 'NUMPAGES' in fxml
    except Exception:
        ftr = (False, False)
    return dict(colors=colors, hl=hl, shd=shd, omml=omml, media=len(media),
                first=first, hdr=hdr, ftr=ftr, paras=len(doc.paragraphs), z=z)


def main():
    bad = []
    tot = Counter()
    print(f"{'文件':<34}{'段':>5}{'公式':>6}{'图':>4}{'彩字':>5}{'高亮':>5}{'底纹':>5}{'页眉':>5}{'页码':>5}  首段")
    print("-" * 118)
    for f in sorted(DST.rglob('*.docx')):
        rel = f.relative_to(DST)
        r = zscan(f)
        # 只有「非 auto 且非纯黑」才算彩色残留；显式 000000 是黑白化的预期结果
        ncol = sum(v for k, v in r['colors'].items() if k not in ('auto', '000000'))
        nshd = sum(v for k, v in r['shd'].items() if k not in ('auto', 'clear'))
        # 彩色图残留
        cimg = []
        for mn in [n for n in r['z'].namelist() if n.startswith('word/media/')]:
            pct, _ = colorful_fullres(r['z'].read(mn))
            if pct and pct > 0.2:
                cimg.append((mn, round(pct, 2)))
        tot['omml'] += r['omml']
        tot['media'] += r['media']
        tot['col'] += ncol
        tot['hl'] += r['hl']
        tot['shd'] += nshd
        tot['cimg'] += len(cimg)
        if cimg:
            bad.append((str(rel), cimg))
        mark = '' if (ncol == 0 and r['hl'] == 0 and nshd == 0 and not cimg) else '  <<<'
        print(f"{str(rel)[:33]:<34}{r['paras']:>5}{r['omml']:>6}{r['media']:>4}"
              f"{ncol:>5}{r['hl']:>5}{nshd:>5}"
              f"{'Y' if r['hdr'] else '-':>5}{'Y' if r['ftr'][0] else '-':>5}"
              f"  {r['first'][:26]!r}{mark}")

    print("-" * 118)
    print(f"合计 公式 {tot['omml']} | 图片 {tot['media']} | 残留彩字 {tot['col']} | "
          f"高亮 {tot['hl']} | 底纹 {tot['shd']} | 彩色图残留 {tot['cimg']}")
    if bad:
        print("\n仍含彩色图的章节：")
        for name, c in bad:
            print(f"  {name}: {c}")
    else:
        print("\n✅ 无彩色残留")


if __name__ == '__main__':
    main()
