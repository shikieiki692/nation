# -*- coding: utf-8 -*-
"""
学生版习题书 → 黑白打印版（发给学生）

处理项：
  1. 删除每章开头的封面块（书名/篇/章/题量/版本/日期）
  2. 黑白化：所有彩色文字→纯黑，去高亮/底纹，边框→黑
  3. 图片：彩色图灰度化+对比度增强（保证黑白打印可辨）
  4. 页眉：篇名 · 章名（散页可归位）
  5. 页脚：第 X 页 / 共 Y 页
  6. 打印友好：标题防孤行、题组段前留白
  ※ 不改字体（用户要求）、不动 OMML 公式与正文文字

用法：
  python build_print_student.py [--dry-run]
"""
import sys, io, re, shutil, zipfile, argparse
from pathlib import Path
from collections import Counter

import docx
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from PIL import Image, ImageOps

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

SRC = Path(r"C:\Obsidion\妙妙屋\00-首页\题组Word\习题书\学生版")
DST = Path(r"C:\Obsidion\妙妙屋\00-首页\题组Word\习题书\学生版-打印版")

STAT = Counter()


# ---------------------------------------------------------------- 工具
def set_keep_next(p):
    """标题与下一段同页，避免孤行"""
    pPr = p._p.get_or_add_pPr()
    kn = OxmlElement('w:keepNext')
    pPr.append(kn)


def build_field(paragraph, parts):
    """parts = [('text', '第 '), ('field', 'PAGE'), ('text', ' 页 / 共 '), ...]"""
    for kind, val in parts:
        if kind == 'text':
            r = paragraph.add_run(val)
            r.font.size = Pt(9)
            r.font.color.rgb = RGBColor(0, 0, 0)
        else:
            r = paragraph.add_run()
            r.font.size = Pt(9)
            r.font.color.rgb = RGBColor(0, 0, 0)
            fld = OxmlElement('w:fldChar')
            fld.set(qn('w:fldCharType'), 'begin')
            r._r.append(fld)

            r2 = paragraph.add_run()
            r2.font.size = Pt(9)
            it = OxmlElement('w:instrText')
            it.set(qn('xml:space'), 'preserve')
            it.text = f' {val} '
            r2._r.append(it)

            r3 = paragraph.add_run()
            r3.font.size = Pt(9)
            sep = OxmlElement('w:fldChar')
            sep.set(qn('w:fldCharType'), 'separate')
            r3._r.append(sep)

            r4 = paragraph.add_run('1')
            r4.font.size = Pt(9)

            r5 = paragraph.add_run()
            r5.font.size = Pt(9)
            end = OxmlElement('w:fldChar')
            end.set(qn('w:fldCharType'), 'end')
            r5._r.append(end)


def set_footer(sec):
    """页脚：第 X 页 / 共 Y 页（居中，9pt 纯黑）"""
    ft = sec.footer
    ft.is_linked_to_previous = False
    # 清空现有内容
    for p in list(ft.paragraphs):
        p._p.getparent().remove(p._p)
    p = ft.add_paragraph()
    p.alignment = 1  # center
    build_field(p, [('text', '第 '), ('field', 'PAGE'),
                    ('text', ' 页 / 共 '), ('field', 'NUMPAGES'), ('text', ' 页')])


def set_header(sec, text):
    """页眉：篇名 · 章名"""
    hd = sec.header
    hd.is_linked_to_previous = False
    for p in list(hd.paragraphs):
        p._p.getparent().remove(p._p)
    p = hd.add_paragraph()
    p.alignment = 2  # right
    r = p.add_run(text)
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0, 0, 0)
    # 页眉下细线
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'), 'single')
    bot.set(qn('w:sz'), '4')
    bot.set(qn('w:space'), '2')
    bot.set(qn('w:color'), '000000')
    pbdr.append(bot)
    pPr.append(pbdr)


def title_of(rel: Path):
    """由路径推出「篇名 · 章名」"""
    fname = rel.stem                       # 1-热力学 / 来源索引 / 附录-...
    parent = rel.parent.name               # 第一篇-化学原理 / 学生版
    if parent and parent != SRC.name and parent.startswith('第'):
        pian = parent.replace('-', ' ', 1)          # 第一篇 化学原理
        m = re.match(r'(\d+)[-、.](.+)', fname)
        zhang = f"第{m.group(1)}章 {m.group(2)}" if m else fname
        return f"{pian} · {zhang}"
    return fname


# ---------------------------------------------------------------- 黑白化（XML 层）
RE_COLOR = re.compile(r'<w:color\b[^>]*/>')
RE_HIGHLIGHT = re.compile(r'<w:highlight\b[^>]*/>')
RE_SHD = re.compile(r'<w:shd\b[^>]*/>')
RE_BORDER = re.compile(r'<(w:top|w:left|w:bottom|w:right|w:insideH|w:insideV|w:tl2br|w:tr2bl)\b([^>]*)/>')
RE_THEME_ATTR = re.compile(r'\s*w:theme(?:Color|Shade|Tint)="[^"]*"')


def fix_color_tag(m):
    tag = m.group(0)
    if 'w:val="auto"' in tag:
        return '<w:color w:val="auto"/>'
    STAT['xml_color'] += 1
    return '<w:color w:val="000000"/>'


def fix_border_tag(m):
    tag = m.group(0)
    STAT['xml_border'] += 1
    tag = RE_THEME_ATTR.sub('', tag)
    tag = re.sub(r'w:color="[^"]*"', 'w:color="000000"', tag)
    return tag


def normalize_xml(xml: str) -> str:
    # 1) 文字颜色 → 纯黑
    xml = RE_COLOR.sub(fix_color_tag, xml)
    # 2) 高亮 → 移除
    n = len(RE_HIGHLIGHT.findall(xml))
    if n:
        STAT['xml_highlight'] += n
        xml = RE_HIGHLIGHT.sub('', xml)
    # 3) 底纹 → 无填充
    def _shd(mm):
        tag = mm.group(0)
        if re.search(r'w:fill="(auto|clear)"', tag):
            return tag
        STAT['xml_shd'] += 1
        return '<w:shd w:val="clear" w:color="auto" w:fill="auto"/>'
    xml = RE_SHD.sub(_shd, xml)
    # 4) 边框颜色 → 黑
    xml = RE_BORDER.sub(fix_border_tag, xml)
    # 5) 主题色残留属性清理
    xml = RE_THEME_ATTR.sub('', xml)
    return xml


# ---------------------------------------------------------------- 图片灰度化
def is_colorful(img: Image.Image, thresh=18, min_ratio=0.2) -> bool:
    """
    全分辨率网格采样检测彩色。
    不使用 resize 缩放——缩放会把小面积彩色（如化学结构式里的彩色原子标记）
    与周边白色平均掉，导致漏检；JPEG 色度噪声同理。
    """
    rgb = img.convert('RGB')
    w, h = rgb.size
    px = rgb.load()
    step = max(1, int((w * h / 60000) ** 0.5))
    tot = col = 0
    for y in range(0, h, step):
        for x in range(0, w, step):
            r, g, b = px[x, y]
            tot += 1
            if max(r, g, b) - min(r, g, b) > thresh:
                col += 1
                # 早停：已超比例阈值
                if col > tot * min_ratio / 100:
                    return True
    return tot and (col / tot * 100) > min_ratio


def to_print_gray(img: Image.Image) -> Image.Image:
    """彩色 → 灰度，并做亮度/对比度校正，保证黑白打印机上可辨"""
    g = img.convert('L')
    # 自动对比度（cutoff 0.5% 避免单像素噪声拉伸）
    g = ImageOps.autocontrast(g, cutoff=(0, 1))
    # 轻度锐化，弥补打印网点扩散
    from PIL import ImageFilter
    g = g.filter(ImageFilter.UnsharpMask(radius=1.2, percent=70, threshold=3))
    return g


def process_media(zin: zipfile.ZipFile, zout, rels_to_write: dict):
    """处理 word/media/*，返回 {name: bytes}"""
    out = {}
    for name in zin.namelist():
        if not name.startswith('word/media/'):
            continue
        data = zin.read(name)
        try:
            img = Image.open(io.BytesIO(data))
            img.load()
        except Exception:
            out[name] = data
            continue
        STAT['img_total'] += 1
        if is_colorful(img):
            STAT['img_colored'] += 1
            g = to_print_gray(img)
            buf = io.BytesIO()
            fmt = (img.format or 'PNG').upper()
            if fmt == 'JPEG':
                g.save(buf, 'JPEG', quality=94, subsampling=0, optimize=True)
            else:
                g.convert('RGB').save(buf, 'PNG', optimize=True)
            out[name] = buf.getvalue()
        else:
            STAT['img_gray_kept'] += 1
            out[name] = data
        rels_to_write[name] = True
    return out


# ---------------------------------------------------------------- 主流程
def convert_one(src: Path, dst: Path, dry=False):
    rel = src.relative_to(SRC)
    label = title_of(rel)

    doc = docx.Document(str(src))

    # ---- 1. 删除封面（首个 Heading 1 之前的全部段落/表格）
    paras = doc.paragraphs
    h1_idx = next((i for i, p in enumerate(paras) if p.style.name == 'Heading 1'), None)
    if h1_idx is None:
        STAT['no_h1'] += 1
        return dict(src=src.name, ok=False, msg='无 Heading 1，跳过封面删除')
    cover_els = [paras[i]._p for i in range(h1_idx)]
    for el in cover_els:
        el.getparent().remove(el)
    STAT['cover_removed'] += len(cover_els)

    # ---- 2. 打印友好间距
    h1 = doc.paragraphs[0]
    h1.paragraph_format.space_before = Pt(0)
    h1.paragraph_format.space_after = Pt(12)
    set_keep_next(h1)

    n_h2 = 0
    for p in doc.paragraphs:
        if p.style.name == 'Heading 2':
            n_h2 += 1
            pf = p.paragraph_format
            before = pf.space_before.pt if pf.space_before else 0
            pf.space_before = Pt(max(before, 14))
            pf.space_after = Pt(4)
            set_keep_next(p)
    STAT['h2_tuned'] += n_h2

    # ---- 3. 页眉 / 页脚
    sec = doc.sections[0]
    set_header(sec, label)
    set_footer(sec)

    if dry:
        return dict(src=str(rel), ok=True, msg=f'cover={len(cover_els)} h2={n_h2} header={label}', dry=True)

    dst.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(dst))

    # ---- 4. 后处理：颜色归一化 + 图片灰度（直接改 zip）
    # 注意：不使用临时文件 + shutil.move —— 移动会触发删除操作，在批量场景下
    # 会被 safe-delete 的批量阈值拦截。改为整体读入内存后直接覆盖写最终文件。
    NORMALIZE_PARTS = ('word/document.xml', 'word/styles.xml',
                       'word/footnotes.xml', 'word/endnotes.xml')

    def _needs_norm(n):
        return n.endswith('.xml') and n.startswith('word/') and (
            n in NORMALIZE_PARTS or 'footer' in n or 'header' in n)

    with zipfile.ZipFile(str(dst)) as zin:
        media_new = process_media(zin, None, {})
        items = [(n, zin.read(n)) for n in zin.namelist()]

    with zipfile.ZipFile(str(dst), 'w', zipfile.ZIP_DEFLATED) as zout:
        for name, data in items:
            data = media_new.get(name, data)
            if _needs_norm(name):
                try:
                    data = normalize_xml(data.decode('utf-8')).encode('utf-8')
                except UnicodeDecodeError:
                    pass
            zout.writestr(name, data)

    return dict(src=str(rel), ok=True, msg=f'cover={len(cover_els)} h2={n_h2}')


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument('--dry-run', action='store_true')
    args = ap.parse_args()

    files = sorted(SRC.rglob('*.docx'))
    print(f"源文件 {len(files)} 个\n目标 {DST}\n" + "-" * 64)

    rows = []
    for f in files:
        rel = f.relative_to(SRC)
        dst = DST / rel
        try:
            r = convert_one(f, dst, dry=args.dry_run)
        except Exception as e:
            r = dict(src=str(rel), ok=False, msg=f'ERR {type(e).__name__}: {e}')
            STAT['fail'] += 1
        rows.append(r)
        flag = 'OK ' if r['ok'] else 'FAIL'
        print(f"[{flag}] {r['src']:<38} {r['msg']}")

    print("-" * 64)
    print(f"完成 {sum(1 for r in rows if r['ok'])}/{len(rows)}")
    print("统计:", dict(STAT))
    if not args.dry_run:
        total = sum(p.stat().st_size for p in DST.rglob('*.docx'))
        print(f"输出 {len(list(DST.rglob('*.docx')))} 个文件，共 {total/1024/1024:.1f} MB")


if __name__ == '__main__':
    main()
