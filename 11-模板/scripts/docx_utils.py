"""
docx_utils.py — 学生讲义 Word 输出共享风格库 (v2)

字体口径（依据用户反馈修正）：
  - 中文正文：宋体 SimSun（备选仿宋 FangSong）
  - 中文标题/表头：黑体 SimHei
  - 英文/数字/变量/普通符号：Times New Roman
  - 仅非常特殊、必须保留中文风格的符号跟随中文字体

本库提供两类使用方式：
  A) 手动构建：base_doc() → title_block() → heading() → para() → simple_table()
  B) Pandoc 后处理：postprocess_pandoc_docx() 修正 pypandoc 生成文档的字体

提取自 tmp/docs/polarity-bridge-docx/build-polarity-bridge-docx.py
          tmp/docs/crystal-structure-docx/build-crystal-structure-docx.py
"""

from copy import deepcopy
from pathlib import Path
from typing import Optional
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

# ── 字体常量 ──────────────────────────────────────────────
CN_BODY_FONT = "SimSun"        # 宋体 — 中文正文 / 表格正文
CN_HEAD_FONT = "SimHei"        # 黑体 — 中文标题 / 表头
EN_FONT = "Times New Roman"    # 英文 / 数字 / 变量 / 一般符号
CN_CAPTION_FONT = "FangSong"   # 仿宋 — 图例题注（仿宋五号居中）

TITLE_COLOR = RGBColor(23, 50, 77)    # 深蓝 — 标题/表头
SUB_COLOR = RGBColor(91, 105, 117)    # 灰蓝 — 副标题/注释

# ── 页面设置 ──────────────────────────────────────────────
PAGE_WIDTH = Cm(21)
PAGE_HEIGHT = Cm(29.7)
MARGIN_LEFT = Cm(2.2)
MARGIN_RIGHT = Cm(2.2)
MARGIN_TOP = Cm(2.0)
MARGIN_BOTTOM = Cm(2.0)

# ── 段落间距常量（pt）──────────────────────────────────────
SPACING_IMAGE_BEFORE = 8      # 图片段落前间距
SPACING_IMAGE_AFTER = 4       # 图片段落后间距
SPACING_TABLE_BEFORE = 6      # 表格前间距
SPACING_TABLE_AFTER = 4       # 表格后间距
SPACING_CAPTION_BEFORE = 2    # 图注与图片间距
SPACING_CAPTION_AFTER = 8     # 图注与正文间距
SPACING_HEADING_BEFORE = 14   # 标题前间距
SPACING_HEADING_AFTER = 6     # 标题后间距
LINE_SPACING_BODY = 1.25      # 正文行距

# ── 字体设置函数 ──────────────────────────────────────────

def set_run_font(
    run,
    size: float = 11.5,
    bold: bool = False,
    color: Optional[RGBColor] = None,
    cn_font: str = CN_BODY_FONT,
    en_font: str = EN_FONT,
) -> None:
    """同时设置中文字体（eastAsia）和西文字体（ascii/hAnsi）。

    约定：
      - 中文内容走 cn_font
      - 英文/数字/一般符号统一走 Times New Roman
      - 仅非常特殊且必须保留中文风格的符号跟随中文字体
    """
    run.font.name = en_font
    rpr = run._element.rPr
    rpr.rFonts.set(qn("w:eastAsia"), cn_font)
    rpr.rFonts.set(qn("w:ascii"), en_font)
    rpr.rFonts.set(qn("w:hAnsi"), en_font)
    rpr.rFonts.set(qn("w:cs"), en_font)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color


def set_paragraph_spacing(
    paragraph,
    after: float = 6,
    before: float = 0,
    line: float = 1.3,
) -> None:
    fmt = paragraph.paragraph_format
    fmt.space_after = Pt(after)
    fmt.space_before = Pt(before)
    fmt.line_spacing = line


def add_text(
    paragraph,
    text: str,
    size: float = 11.5,
    bold: bool = False,
    color: Optional[RGBColor] = None,
    cn_font: str = CN_BODY_FONT,
) -> None:
    """给段落追加一个 run，自动设置字体。"""
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color, cn_font=cn_font)


# ── 手动构建函数（用于非 pandoc 场景）─────────────────────

def base_doc() -> Document:
    """A4 竖排页面，标准页边距，Normal 默认为 SimSun + TNR。"""
    doc = Document()
    section = doc.sections[0]
    section.page_width = PAGE_WIDTH
    section.page_height = PAGE_HEIGHT
    section.left_margin = MARGIN_LEFT
    section.right_margin = MARGIN_RIGHT
    section.top_margin = MARGIN_TOP
    section.bottom_margin = MARGIN_BOTTOM

    normal = doc.styles["Normal"]
    normal.font.name = EN_FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), CN_BODY_FONT)
    normal._element.rPr.rFonts.set(qn("w:ascii"), EN_FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), EN_FONT)
    normal._element.rPr.rFonts.set(qn("w:cs"), EN_FONT)
    normal.font.size = Pt(11.5)
    normal.paragraph_format.line_spacing = LINE_SPACING_BODY
    return doc


def title_block(doc: Document, title: str, subtitle: Optional[str] = None) -> None:
    """居中标题块：大标题（SimHei 24pt）+ 副标题（灰色 11pt）。"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, after=3, line=1.15)
    add_text(p, title, size=24, bold=True, color=TITLE_COLOR, cn_font=CN_HEAD_FONT)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p2, after=12, line=1.15)
    if subtitle:
        add_text(p2, subtitle, size=11, color=SUB_COLOR, cn_font=CN_BODY_FONT)
    else:
        add_text(p2, "", size=11)


def heading(doc: Document, text: str, level: int = 1) -> None:
    """标题段落（SimHei，level 1=15pt，level 2=12.5pt）。"""
    sz = 15 if level == 1 else 12.5
    before = 8 if level == 1 else 4
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=before, after=4, line=1.2)
    add_text(p, text, size=sz, bold=True, color=TITLE_COLOR, cn_font=CN_HEAD_FONT)
    return p


def para(doc: Document, text: str) -> None:
    """正文段落（SimSun 11.5pt）。"""
    p = doc.add_paragraph()
    set_paragraph_spacing(p, after=5)
    add_text(p, text, size=11.5, cn_font=CN_BODY_FONT)
    return p


def bullet(doc: Document, text: str, level: int = 0) -> None:
    """项目符号段落（SimSun 11.5pt）。"""
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(0.74 + 0.5 * level)
    p.paragraph_format.first_line_indent = Cm(-0.45)
    set_paragraph_spacing(p, after=3)
    add_text(p, text, size=11.5, cn_font=CN_BODY_FONT)
    return p


def simple_table(doc: Document, rows: list[list[str]],
                 widths: Optional[list] = None, font_size: float = 10.5) -> None:
    """简单表格：表头自动黑体加粗，正文自动宋体，内容居中。"""
    if not rows:
        return
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    table.autofit = False
    if widths is None:
        widths_cm = [Cm(3.2)] * len(rows[0])
    else:
        widths_cm = widths
    for r, row in enumerate(rows):
        for c, value in enumerate(row):
            cell = table.cell(r, c)
            cell.width = widths_cm[c]
            cell_p = cell.paragraphs[0]
            cell_p.alignment = WD_ALIGN_PARAGRAPH.CENTER  # 内容居中
            set_paragraph_spacing(cell_p, after=0, line=1.15)
            add_text(cell_p, str(value), size=font_size,
                     bold=(r == 0), color=TITLE_COLOR if r == 0 else None,
                     cn_font=CN_HEAD_FONT if r == 0 else CN_BODY_FONT)
    doc.add_paragraph()


def picture_block(doc: Document, image_path: str, caption: str = "",
                  width_cm: float = 12.0) -> None:
    """图片 + 居中题注（仿宋五号）。"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, after=2, before=2, line=1.0)
    run = p.add_run()
    run.add_picture(str(image_path), width=Cm(width_cm))

    if caption:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_spacing(cap, after=6, line=1.0)
        add_text(cap, caption, size=10.5, color=SUB_COLOR, cn_font=CN_CAPTION_FONT)


# ── OMML 公式字号辅助函数（v3 新增）────────────────────────

def _set_display_math_font_size(paragraph, size_pt: float = 14) -> bool:
    """将段落中 OMML 显示公式（m:oMathPara）的默认字号设为 size_pt。

    Pandoc 将 $$\\dots$$ 显示公式转换为 m:oMathPara 包装的 OMML 结构，
    默认继承段落正文的 11.5pt。本函数在 m:ctrlPr 中插入 w:sz/w:szCs，
    使公式以 size_pt（四号 = 14pt）渲染。

    返回 True 表示已处理，False 表示该段落无显示公式。
    """
    para_elem = paragraph._element
    oMathPara = para_elem.find(qn("m:oMathPara"))
    if oMathPara is None:
        # Also check inline math (m:oMath without oMathPara wrapper)
        # Only resize if formula text content > 5 chars
        oMath_list = para_elem.findall(qn("m:oMath"))
        handled = False
        for oMath in oMath_list:
            # Skip if this oMath is inside an oMathPara (already handled above)
            if oMath.getparent() is not None and oMath.getparent().tag == qn("m:oMathPara"):
                continue
            # Measure text content of the formula
            text_parts = []
            for t_elem in oMath.iter(qn("m:t")):
                if t_elem.text:
                    text_parts.append(t_elem.text)
            formula_text = "".join(text_parts)
            # Count non-whitespace characters
            char_count = sum(1 for c in formula_text if not c.isspace())
            if char_count > 5:
                _inject_ctrlpr_font_size(oMath, qn("m:oMathPr"), "m:oMathPr", size_pt)
                handled = True
        return handled

    _inject_ctrlpr_font_size(oMathPara, qn("m:oMathParaPr"), "m:oMathParaPr", size_pt)
    return True


def _inject_ctrlpr_font_size(
    parent_elem,
    pr_qn_tag: str,      # namespace-qualified tag for find()
    pr_prefix_tag: str,   # prefix:localname tag for OxmlElement()
    size_pt: float,
) -> None:
    """在 parent_elem（oMathPara 或 oMath）的 pr 子元素中插入字号设置。"""
    pr = parent_elem.find(pr_qn_tag)
    if pr is None:
        pr = OxmlElement(pr_prefix_tag)
        parent_elem.insert(0, pr)

    ctrlPr = pr.find(qn("m:ctrlPr"))
    if ctrlPr is None:
        ctrlPr = OxmlElement("m:ctrlPr")
        pr.append(ctrlPr)

    rPr = ctrlPr.find(qn("w:rPr"))
    if rPr is None:
        rPr = OxmlElement("w:rPr")
        ctrlPr.append(rPr)

    half_pts = str(int(size_pt * 2))
    sz = rPr.find(qn("w:sz"))
    if sz is None:
        sz = OxmlElement("w:sz")
        rPr.append(sz)
    sz.set(qn("w:val"), half_pts)

    szCs = rPr.find(qn("w:szCs"))
    if szCs is None:
        szCs = OxmlElement("w:szCs")
        rPr.append(szCs)
    szCs.set(qn("w:val"), half_pts)


def _is_cjk_char(ch: str) -> bool:
    """Return True when the character is a CJK ideograph."""
    cp = ord(ch)
    return (
        0x4E00 <= cp <= 0x9FFF or    # CJK统一汉字
        0x3400 <= cp <= 0x4DBF or    # CJK扩展A
        0xF900 <= cp <= 0xFAFF       # CJK兼容汉字
    )


def _has_cjk(text: str) -> bool:
    """检测文本是否包含CJK字符（中日韩统一表意文字）。"""
    return any(_is_cjk_char(ch) for ch in text)


SUPERSCRIPT_CHAR_MAP = {
    "⁰": "0",
    "¹": "1",
    "²": "2",
    "³": "3",
    "⁴": "4",
    "⁵": "5",
    "⁶": "6",
    "⁷": "7",
    "⁸": "8",
    "⁹": "9",
    "⁺": "+",
    "⁻": "-",
    "⁼": "=",
    "⁽": "(",
    "⁾": ")",
}

SUBSCRIPT_CHAR_MAP = {
    "₀": "0",
    "₁": "1",
    "₂": "2",
    "₃": "3",
    "₄": "4",
    "₅": "5",
    "₆": "6",
    "₇": "7",
    "₈": "8",
    "₉": "9",
    "₊": "+",
    "₋": "-",
    "₌": "=",
    "₍": "(",
    "₎": ")",
}


def _unicode_script_role(ch: str) -> Optional[str]:
    """Return the vertical alignment role for a unicode script char."""
    if ch in SUPERSCRIPT_CHAR_MAP:
        return "superscript"
    if ch in SUBSCRIPT_CHAR_MAP:
        return "subscript"
    return None


def _normalize_unicode_script_char(ch: str) -> str:
    """Map unicode super/subscript glyphs to their plain-text equivalent."""
    if ch in SUPERSCRIPT_CHAR_MAP:
        return SUPERSCRIPT_CHAR_MAP[ch]
    if ch in SUBSCRIPT_CHAR_MAP:
        return SUBSCRIPT_CHAR_MAP[ch]
    return ch


def _has_unicode_script_chars(text: str) -> bool:
    """Check whether text contains unicode super/subscript glyphs."""
    return any(_unicode_script_role(ch) is not None for ch in text)


def _segment_text_by_script(text: str) -> list[tuple[str, bool]]:
    """Split text into CJK and non-CJK segments.

    The returned boolean indicates whether the segment should use the CJK font.
    Symbols such as `²⁺⁻` are treated as non-CJK and should therefore follow
    the western font policy (Times New Roman).
    """
    if not text:
        return []
    segments: list[tuple[str, bool]] = []
    current_chars = [text[0]]
    current_is_cjk = _is_cjk_char(text[0])
    for ch in text[1:]:
        is_cjk = _is_cjk_char(ch)
        if is_cjk == current_is_cjk:
            current_chars.append(ch)
            continue
        segments.append(("".join(current_chars), current_is_cjk))
        current_chars = [ch]
        current_is_cjk = is_cjk
    segments.append(("".join(current_chars), current_is_cjk))
    return segments


def _run_is_safe_to_split(run) -> bool:
    """Only split plain-text runs; skip drawings, fields, and other OOXML nodes."""
    allowed = {
        qn("w:rPr"),
        qn("w:t"),
        qn("w:tab"),
        qn("w:br"),
        qn("w:cr"),
        qn("w:noBreakHyphen"),
        qn("w:softHyphen"),
    }
    return all(child.tag in allowed for child in run._element)


def _copy_run_rpr(src_run, dst_run) -> None:
    """Copy run properties (rPr) from src to dst."""
    dst_rpr = dst_run._element.rPr
    if dst_rpr is not None:
        dst_run._element.remove(dst_rpr)
    src_rpr = src_run._element.rPr
    if src_rpr is not None:
        dst_run._element.insert(0, deepcopy(src_rpr))


def _segment_text_by_vertical_align(text: str) -> list[tuple[str, Optional[str]]]:
    """Split text into baseline / superscript / subscript segments."""
    if not text:
        return []
    segments: list[tuple[str, Optional[str]]] = []
    current_role = _unicode_script_role(text[0])
    current_chars = [_normalize_unicode_script_char(text[0])]
    for ch in text[1:]:
        role = _unicode_script_role(ch)
        mapped = _normalize_unicode_script_char(ch)
        if role == current_role:
            current_chars.append(mapped)
            continue
        segments.append(("".join(current_chars), current_role))
        current_chars = [mapped]
        current_role = role
    segments.append(("".join(current_chars), current_role))
    return segments


def _apply_vertical_align(run, role: Optional[str]) -> None:
    """Apply Word superscript/subscript formatting to a run."""
    run.font.superscript = role == "superscript"
    run.font.subscript = role == "subscript"


def _split_mixed_script_runs_in_paragraph(paragraph) -> None:
    """Split runs that mix CJK with latin/symbol text so fonts can be assigned precisely."""
    original_runs = list(paragraph.runs)
    for run in original_runs:
        text = run.text or ""
        if not text or not _has_cjk(text):
            continue
        if not _run_is_safe_to_split(run):
            continue
        segments = _segment_text_by_script(text)
        if len(segments) <= 1:
            continue
        run.text = segments[0][0]
        prev_run = run
        for seg_text, _ in segments[1:]:
            new_run = paragraph.add_run(seg_text)
            _copy_run_rpr(run, new_run)
            prev_run._element.addnext(new_run._element)
            prev_run = new_run


def _split_unicode_script_runs_in_paragraph(paragraph) -> None:
    """Convert unicode super/subscript glyphs into dedicated Word-format runs."""
    original_runs = list(paragraph.runs)
    for run in original_runs:
        text = run.text or ""
        if not text or not _has_unicode_script_chars(text):
            continue
        if not _run_is_safe_to_split(run):
            continue
        segments = _segment_text_by_vertical_align(text)
        if not segments:
            continue

        first_text, first_role = segments[0]
        run.text = first_text
        _apply_vertical_align(run, first_role)

        prev_run = run
        for seg_text, seg_role in segments[1:]:
            new_run = paragraph.add_run(seg_text)
            _copy_run_rpr(run, new_run)
            _apply_vertical_align(new_run, seg_role)
            prev_run._element.addnext(new_run._element)
            prev_run = new_run


def _iter_all_paragraphs(doc: Document):
    """Yield paragraphs in body and tables."""
    for para in doc.paragraphs:
        yield para

    def _iter_tables(tables):
        for table in tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        yield para
                    yield from _iter_tables(cell.tables)

    yield from _iter_tables(doc.tables)


def _split_mixed_script_runs(doc: Document) -> None:
    """Pre-split mixed-script runs before font normalization."""
    for para in _iter_all_paragraphs(doc):
        _split_mixed_script_runs_in_paragraph(para)


def _split_unicode_script_runs(doc: Document) -> None:
    """Convert unicode super/subscript glyphs before font normalization."""
    for para in _iter_all_paragraphs(doc):
        _split_unicode_script_runs_in_paragraph(para)


def _resize_and_center_images(doc: Document) -> None:
    """约束超大图片宽度、根据宽高比自适应缩放、居中。

    策略：
    1. 根据图片原始宽高比自适应缩放：
       - 横向图（宽>高）：最大 10cm 宽
       - 纵向图（高>宽）：最大 7cm 宽
       - 方形图：最大 9cm 宽
       - 超小图（<4cm原宽）：放大到 5cm
    2. 所有图片保持 inline 模式，居中对齐

    文本区宽度 = A4(21cm) - 左2.2cm - 右2.2cm = 16.6cm。
    """
    W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    WP_NS = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
    A_NS = "http://schemas.openxmlformats.org/drawingml/2006/main"

    # EMU constants
    CM_TO_EMU = 360000  # 1cm = 360000 EMU
    MAX_LANDSCAPE = int(10.0 * CM_TO_EMU)  # 横向图最大 10cm
    MAX_PORTRAIT = int(7.0 * CM_TO_EMU)    # 纵向图最大 7cm
    MAX_SQUARE = int(9.0 * CM_TO_EMU)      # 方形图最大 9cm
    MIN_WIDTH = int(4.0 * CM_TO_EMU)       # 小于此宽度的图放大到 5cm
    ENLARGE_TARGET = int(5.0 * CM_TO_EMU)

    body = doc.element.body

    for p_elem in body.iter(f"{{{W_NS}}}p"):
        drawings = p_elem.findall(f".//{{{W_NS}}}drawing")
        if not drawings:
            continue

        for drawing in drawings:
            inline = drawing.find(f"{{{WP_NS}}}inline")
            if inline is None:
                continue
            extent = inline.find(f"{{{WP_NS}}}extent")
            if extent is None:
                continue
            cx_str = extent.get("cx")
            cy_str = extent.get("cy")
            if not cx_str or not cy_str:
                continue
            cx, cy = int(cx_str), int(cy_str)
            if cx <= 0 or cy <= 0:
                continue

            # 根据宽高比确定最大宽度
            aspect = cx / cy
            if aspect > 1.3:  # 横向图
                max_w = MAX_LANDSCAPE
            elif aspect < 0.7:  # 纵向图
                max_w = MAX_PORTRAIT
            else:  # 方形图
                max_w = MAX_SQUARE

            # 超小图放大
            if cx < MIN_WIDTH:
                max_w = ENLARGE_TARGET

            # 缩放
            new_cx = cx
            new_cy = cy
            if cx > max_w:
                ratio = max_w / cx
                new_cx = max_w
                new_cy = int(cy * ratio)
            elif cx < MIN_WIDTH:
                ratio = ENLARGE_TARGET / cx
                new_cx = ENLARGE_TARGET
                new_cy = int(cy * ratio)

            # 更新 extent
            extent.set("cx", str(new_cx))
            extent.set("cy", str(new_cy))
            # 同步更新 a:ext
            for aext in drawing.iter(f"{{{A_NS}}}ext"):
                try:
                    cur_cx = int(aext.get("cx", "0"))
                    cur_cy = int(aext.get("cy", "0"))
                except (ValueError, TypeError):
                    continue
                if cur_cx > 0 and cur_cy > 0 and abs(cur_cx / cur_cy - cx / cy) < 0.1:
                    aext.set("cx", str(new_cx))
                    aext.set("cy", str(new_cy))
                    break

        # 居中段落
        pPr = p_elem.find(f"{{{W_NS}}}pPr")
        if pPr is None:
            pPr = OxmlElement("w:pPr")
            p_elem.insert(0, pPr)
        jc = pPr.find(f"{{{W_NS}}}jc")
        if jc is None:
            jc = OxmlElement("w:jc")
            pPr.append(jc)
        jc.set(f"{{{W_NS}}}val", "center")


def _convert_inline_to_wrap_top_bottom(drawing, inline, cx, cy):
    """将 wp:inline 转为 wp:anchor + wrapTopAndBottom。

    wrapTopAndBottom 让图片浮动，文字只在图片上方和下方流动，
    不会在两侧，适合学术文档的居中图片布局。
    """
    WP_NS = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"

    # 创建 anchor 元素
    anchor = OxmlElement("wp:anchor")
    anchor.set("distT", "0")
    anchor.set("distB", "0")
    anchor.set("distL", "114300")   # 1mm 左侧距离
    anchor.set("distR", "114300")   # 1mm 右侧距离
    anchor.set("simplePos", "0")
    anchor.set("relativeHeight", "251658240")
    anchor.set("behindDoc", "0")
    anchor.set("locked", "0")
    anchor.set("layoutInCell", "1")
    anchor.set("allowOverlap", "1")

    # 简单位置（居中：相对于页面水平居中）
    simplePos = OxmlElement("wp:simplePos")
    simplePos.set("x", "0")
    simplePos.set("y", "0")
    anchor.append(simplePos)

    # 水平位置：相对于栏居中
    positionH = OxmlElement("wp:positionH")
    positionH.set("relativeFrom", "column")
    posOffset = OxmlElement("wp:posOffset")
    posOffset.text = "0"  # 居中
    positionH.append(posOffset)
    anchor.append(positionH)

    # 垂直位置：相对于段落
    positionV = OxmlElement("wp:positionV")
    positionV.set("relativeFrom", "paragraph")
    posOffset2 = OxmlElement("wp:posOffset")
    posOffset2.text = "0"
    positionV.append(posOffset2)
    anchor.append(positionV)

    # extent（尺寸）
    extent = OxmlElement("wp:extent")
    extent.set("cx", str(cx))
    extent.set("cy", str(cy))
    anchor.append(extent)

    # effectExtent
    effectExtent = OxmlElement("wp:effectExtent")
    effectExtent.set("l", "0")
    effectExtent.set("t", "0")
    effectExtent.set("r", "0")
    effectExtent.set("b", "0")
    anchor.append(effectExtent)

    # wrapTopAndBottom — 关键！
    wrapTab = OxmlElement("wp:wrapTopAndBottom")
    anchor.append(wrapTab)

    # docPr
    docPr = inline.find(f"{{{WP_NS}}}docPr")
    if docPr is not None:
        anchor.append(docPr)

    # graphic（图片数据）
    graphic = inline.find(f"{{{WP_NS}}}graphic")
    if graphic is not None:
        anchor.append(graphic)

    # 替换 inline 为 anchor
    drawing.remove(inline)
    drawing.insert(0, anchor)


def _has_image(para) -> bool:
    """判断段落是否包含图片（w:drawing 或 w:pict）。"""
    el = para._element
    return bool(el.findall('.//' + qn('w:drawing')) or
                el.findall('.//' + qn('w:pict')))


def _center_first_heading(doc: Document) -> None:
    """Center the first Heading 1 paragraph used as the handout title."""
    for para in doc.paragraphs:
        pPr = para._element.find(qn("w:pPr"))
        if pPr is None:
            continue
        pStyle = pPr.find(qn("w:pStyle"))
        if pStyle is None:
            continue
        if pStyle.get(qn("w:val")) == "Heading1":
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            return


def _paragraph_has_page_field(paragraph) -> bool:
    """Check whether a footer paragraph already contains a PAGE field."""
    for instr in paragraph._element.iter(qn("w:instrText")):
        if (instr.text or "").strip() == "PAGE":
            return True
    return False


def _add_page_numbers(doc: Document) -> None:
    """Insert a centered PAGE field into every section footer."""
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if _paragraph_has_page_field(p):
            continue

        run = p.add_run()
        set_run_font(run, size=10, cn_font=EN_FONT, en_font=EN_FONT)
        fld_char_begin = OxmlElement("w:fldChar")
        fld_char_begin.set(qn("w:fldCharType"), "begin")
        run._element.append(fld_char_begin)

        run2 = p.add_run()
        set_run_font(run2, size=10, cn_font=EN_FONT, en_font=EN_FONT)
        instr_text = OxmlElement("w:instrText")
        instr_text.set(qn("xml:space"), "preserve")
        instr_text.text = " PAGE "
        run2._element.append(instr_text)

        run3 = p.add_run()
        set_run_font(run3, size=10, cn_font=EN_FONT, en_font=EN_FONT)
        fld_char_end = OxmlElement("w:fldChar")
        fld_char_end.set(qn("w:fldCharType"), "end")
        run3._element.append(fld_char_end)


def postprocess_pandoc_docx(
    input_path: Path,
    output_path: Optional[Path] = None,
    body_font: str = CN_BODY_FONT,
    head_font: str = CN_HEAD_FONT,
    center_title: bool = True,
    add_page_numbers: bool = True,
) -> Document:
    """修正 pandoc 生成 docx 的字体和样式。

    参数
    ----------
    input_path : Path
        Pandoc 输出的 .docx 文件路径。
    output_path : Path or None
        修正后保存路径。为 None 时覆盖原文件。
    body_font : str
        中文正文字体（默认 SimSun）。
    head_font : str
        中文标题/表头字体（默认 SimHei）。
    center_title : bool
        是否将首个 Heading 1 题头居中。
    add_page_numbers : bool
        是否在页脚写入 PAGE 域。

    返回
    -------
    docx.Document
    """
    doc = Document(str(input_path))

    # 先按脚本拆分混合 run，避免 `Fe²⁺配位` 这类文本整段继承中文字体。
    _split_mixed_script_runs(doc)
    # 再把 `Fe²⁺ / H₂O / SO₄²⁻` 这类 Unicode 上下标改成 Word 原生上下标。
    _split_unicode_script_runs(doc)

    # ── 全库样式字体清理：遍历所有样式，移除主题引用 ──
    for style in doc.styles:
        try:
            sid = style.style_id
        except AttributeError:
            continue
        # 先确保 rPr 存在（font.name 的 setter 会创建它）
        try:
            style.font.name = EN_FONT
        except Exception:
            continue

        rPr = style._element.rPr
        if rPr is None:
            continue
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            continue

        # 判断类别：标题类用 head_font，正文类用 body_font
        is_heading = bool(
            re.match(r"^(heading|Heading|Heading\s*\d+)", sid)
            or sid in ("Title", "Subtitle", "TOCHeading")
        )
        cn = head_font if is_heading else body_font

        # 补 eastAsia（若缺失）
        if not rFonts.get(qn("w:eastAsia")):
            rFonts.set(qn("w:eastAsia"), cn)
        # 补 ascii/hAnsi（若缺失）
        if not rFonts.get(qn("w:ascii")):
            rFonts.set(qn("w:ascii"), EN_FONT)
        if not rFonts.get(qn("w:hAnsi")):
            rFonts.set(qn("w:hAnsi"), EN_FONT)
        if not rFonts.get(qn("w:cs")):
            rFonts.set(qn("w:cs"), EN_FONT)

        # 移除全部主题引用 → 等线（DengXian）根源
        for _attr in (
            "w:eastAsiaTheme", "w:asciiTheme",
            "w:hAnsiTheme", "w:cstheme",
        ):
            key = qn(_attr)
            if key in rFonts.attrib:
                del rFonts.attrib[key]

    # ── 全库段落级 run 字体清理 ──
    for para in doc.paragraphs:
        # 判断段落属于标题还是正文
        pPr = para._element.find(qn("w:pPr"))
        if pPr is None:
            continue
        pStyle = pPr.find(qn("w:pStyle"))
        sv = pStyle.get(qn("w:val"), "") if pStyle is not None else ""
        is_heading = bool(
            sv.startswith("Heading") or sv in ("Title", "Subtitle", "TOCHeading")
        )
        cn = head_font if is_heading else body_font

        for run in para.runs:
            run.font.name = EN_FONT
            rPr = run._element.rPr
            if rPr is None:
                continue
            rFonts = rPr.find(qn("w:rFonts"))
            if rFonts is None:
                continue
            has_theme = any("Theme" in str(k) for k in rFonts.attrib)
            # 原逻辑：无主题且已有 eastAsia 字体 → 跳过
            # 新逻辑：非CJK run 强制TNR（Unicode亚脚字符需要TNR渲染）
            run_text = run.text or ""
            if not has_theme and rFonts.get(qn("w:eastAsia")) and _has_cjk(run_text):
                continue
            # 选择中文字体：非CJK内容用TNR，CJK内容用原设定
            cn = EN_FONT if not _has_cjk(run_text) else (head_font if is_heading else body_font)
            rFonts.set(qn("w:eastAsia"), cn)
            if not rFonts.get(qn("w:ascii")):
                rFonts.set(qn("w:ascii"), EN_FONT)
            if not rFonts.get(qn("w:hAnsi")):
                rFonts.set(qn("w:hAnsi"), EN_FONT)
            if not rFonts.get(qn("w:cs")):
                rFonts.set(qn("w:cs"), EN_FONT)
            for _attr in (
                "w:eastAsiaTheme", "w:asciiTheme",
                "w:hAnsiTheme", "w:cstheme",
            ):
                key = qn(_attr)
                if key in rFonts.attrib:
                    del rFonts.attrib[key]

    # ── 表格字体 + 内容居中 ──
    for table in doc.tables:
        for i, row in enumerate(table.rows):
            cf = head_font if i == 0 else body_font
            for cell in row.cells:
                for para in cell.paragraphs:
                    # 表格内容居中
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in para.runs:
                        run.font.name = EN_FONT
                        rPr = run._element.rPr
                        if rPr is None:
                            continue
                        rFonts = rPr.find(qn("w:rFonts"))
                        if rFonts is None:
                            continue
                        # 表格中非CJK内容也用TNR（Unicode亚脚字符等）
                        run_text = run.text or ""
                        cf_run = EN_FONT if not _has_cjk(run_text) else cf
                        rFonts.set(qn("w:eastAsia"), cf_run)
                        rFonts.set(qn("w:ascii"), EN_FONT)
                        rFonts.set(qn("w:hAnsi"), EN_FONT)
                        rFonts.set(qn("w:cs"), EN_FONT)
                        for _attr in (
                            "w:eastAsiaTheme", "w:asciiTheme",
                            "w:hAnsiTheme", "w:cstheme",
                        ):
                            key = qn(_attr)
                            if key in rFonts.attrib:
                                del rFonts.attrib[key]
                        if i == 0:
                            run.bold = True

    # ── 表格前后间距：表格前段落加 space_after，表格后段落加 space_before ──
    body_el = doc.element.body
    children = list(body_el)
    for idx, child in enumerate(children):
        if child.tag == qn('w:tbl'):
            # 表格前的段落 → 加 space_after
            if idx > 0 and children[idx - 1].tag == qn('w:p'):
                for p in doc.paragraphs:
                    if p._element is children[idx - 1]:
                        p.paragraph_format.space_after = Pt(SPACING_TABLE_AFTER)
                        break
            # 表格后的段落 → 加 space_before
            if idx + 1 < len(children) and children[idx + 1].tag == qn('w:p'):
                for p in doc.paragraphs:
                    if p._element is children[idx + 1]:
                        p.paragraph_format.space_before = Pt(SPACING_TABLE_BEFORE)
                        break

    # ── 图例题注：识别以"图"/"表"开头的居中短段落 → 仿宋五号 ──
    #    Pandoc 将 ![*caption*](image.png) 输出为居中图片+居中斜体段落，
    #    这些斜体段落通常格式为：纯文本、居中、以"图 "或"表 "开头。
    for para in doc.paragraphs:
        text_content = para.text.strip()
        if not text_content:
            continue
        if text_content.startswith("图 ") or text_content.startswith("表 "):
            # 强制居中（不论 pandoc 原始对齐方式）
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            # 图注间距：与上方图片紧凑，与下方正文留呼吸感
            para.paragraph_format.space_before = Pt(SPACING_CAPTION_BEFORE)
            para.paragraph_format.space_after = Pt(SPACING_CAPTION_AFTER)
            for run in para.runs:
                run.font.size = Pt(10.5)  # 五号
                run.font.bold = False
                run.font.italic = False  # 去斜体
                rPr = run._element.rPr
                if rPr is None:
                    continue
                rFonts = rPr.find(qn("w:rFonts"))
                if rFonts is None:
                    continue
                rFonts.set(qn("w:eastAsia"), CN_CAPTION_FONT)
                rFonts.set(qn("w:ascii"), EN_FONT)
                rFonts.set(qn("w:hAnsi"), EN_FONT)
                rFonts.set(qn("w:cs"), EN_FONT)

    # ── OMML 公式字号：显示公式（$$...$$）→ 四号（14pt）──
    #    长行内公式（>5 非空字符）→ 同样 14pt
    for para in doc.paragraphs:
        _set_display_math_font_size(para, size_pt=14)

    # ── 教师注解着色：教学洞察 / 课堂原话 / 易错点 → 深蓝色 ──
    TEACHER_LABELS = ('教学洞察', '课堂原话', '易错点', '注意', '高频错误',
                      '理解要点', '记忆', '竞赛启示', '练一练', '算一算',
                      '数值冲击', '核心思想', '关键认识', '掌握性要求', '应用')
    TEACHER_COLOR = RGBColor(30, 64, 175)  # #1e40af 深蓝
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        if any(label in text for label in TEACHER_LABELS):
            for run in para.runs:
                run.font.color.rgb = TEACHER_COLOR

    # ── 标题段落间距：Heading 前后增加间距，避免紧贴上文 ──
    for para in doc.paragraphs:
        style_name = para.style.name or ""
        if style_name.startswith("Heading"):
            para.paragraph_format.space_before = Pt(SPACING_HEADING_BEFORE)
            para.paragraph_format.space_after = Pt(SPACING_HEADING_AFTER)

    # ── 图片段落间距：图片前后增加间距，呼吸感 ──
    for para in doc.paragraphs:
        if _has_image(para):
            para.paragraph_format.space_before = Pt(SPACING_IMAGE_BEFORE)
            para.paragraph_format.space_after = Pt(SPACING_IMAGE_AFTER)

    # ── 图片约束与居中：防止宽图溢出页边距 ──
    _resize_and_center_images(doc)

    # ── 页面设置 ──
    section = doc.sections[0]
    section.page_width = PAGE_WIDTH
    section.page_height = PAGE_HEIGHT
    section.left_margin = MARGIN_LEFT
    section.right_margin = MARGIN_RIGHT
    section.top_margin = MARGIN_TOP
    section.bottom_margin = MARGIN_BOTTOM

    if center_title:
        _center_first_heading(doc)
    if add_page_numbers:
        _add_page_numbers(doc)

    # 保存
    save_path = output_path or input_path
    doc.save(str(save_path))
    return doc
