import re, copy
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Mm, Emu
from lxml import etree

BACKUP = r"C:/Users/蕾赛/Desktop/第一轮-4.5小时综合测试（A卷+B卷）（备份）.docx"
OUT    = r"C:/Users/蕾赛/Desktop/第一轮-4.5小时综合测试（A卷+B卷）.docx"
TNR = "Times New Roman"
BLANK_FACTOR = 3  # each ＿ -> 3 underlined spaces (longer than before)

SRC_PAT = re.compile(
    r'[（(]'
    r'(?:20\d\d[^（）]*?(?:卷|预赛)[^（）]*?题[^（）]*?'
    r'|赵鑫光·[^（）]*?'
    r'|汇智·[^（）]*?'
    r'|上海中学竞赛课程·[^（）]*?'
    r'|初赛讲义·[^（）]*?'
    r'|《结构化学基础》习题解析·[^（）]*?'
    r'|省预赛·[^（）]*?'
    r'|第\d+届初赛·[^（）]*?)'
    r'[）)]'
)

def is_cjk(ch):
    o = ord(ch)
    return (0x4E00 <= o <= 0x9FFF or 0x3400 <= o <= 0x4DBF or
            0xF900 <= o <= 0xFAFF or 0x3000 <= o <= 0x303F or
            0xFF00 <= o <= 0xFFEF or 0x2000 <= o <= 0x206F)

XMLSPACE = "{http://www.w3.org/XML/1998/namespace}space"

d = Document(BACKUP)

def ensure_rPr(el):
    rPr = el.find(qn('w:rPr'))
    if rPr is None:
        rPr = etree.SubElement(el, qn('w:rPr'))
    return rPr

def set_text(el, text):
    for wt in el.findall(qn('w:t')):
        el.remove(wt)
    if text == '':
        return
    t = etree.SubElement(el, qn('w:t'))
    t.set(XMLSPACE, 'preserve')
    t.text = text

def set_underline(el, on):
    rPr = ensure_rPr(el)
    u = rPr.find(qn('w:u'))
    if u is None:
        u = etree.SubElement(rPr, qn('w:u'))
    u.set(qn('w:val'), 'single' if on else 'none')

def remove_hint(el):
    rPr = el.find(qn('w:rPr'))
    if rPr is not None:
        rf = rPr.find(qn('w:rFonts'))
        if rf is not None and qn('w:hint') in rf.attrib:
            del rf.attrib[qn('w:hint')]

def set_tnr(el):
    rPr = ensure_rPr(el)
    rf = rPr.find(qn('w:rFonts'))
    if rf is None:
        rf = etree.SubElement(rPr, qn('w:rFonts'))
    txt = ''.join((wt.text or '') for wt in el.findall(qn('w:t')))
    has_cjk = any(is_cjk(c) for c in txt)
    rf.set(qn('w:ascii'), TNR)
    rf.set(qn('w:hAnsi'), TNR)
    rf.set(qn('w:cs'), TNR)
    if not has_cjk:
        rf.set(qn('w:eastAsia'), TNR)

def tokenize(text):
    tokens = []
    buf = ''
    buf_blank = None
    for ch in text:
        if ch == '★':
            continue
        is_blank = (ch == '＿')
        ch2 = (' ' * BLANK_FACTOR) if is_blank else ch
        if buf_blank is None:
            buf_blank = is_blank
            buf = ch2
        elif is_blank == buf_blank:
            buf += ch2
        else:
            tokens.append((buf_blank, buf))
            buf_blank = is_blank
            buf = ch2
    if buf != '':
        tokens.append((buf_blank, buf))
    return tokens

def process_run(el):
    # skip runs without text (images/drawings, fields, empty) — never delete those
    if el.find(qn('w:t')) is None:
        return
    text = ''.join((wt.text or '') for wt in el.findall(qn('w:t')))
    text = SRC_PAT.sub('', text)          # strip source citation
    tokens = tokenize(text)               # drop ★, ＿->underlined spaces
    if not tokens:
        parent = el.getparent()
        if parent is not None:
            parent.remove(el)
        return
    prev = None
    first = True
    for is_blank, chars in tokens:
        if first:
            nr = el
            first = False
        else:
            nr = copy.deepcopy(el)
            (prev if prev is not None else el).addnext(nr)
        set_text(nr, chars)
        set_underline(nr, is_blank)
        set_tnr(nr)
        remove_hint(nr)
        prev = nr

# ---- paragraph-level source removal (handles parens split across runs) ----
def clean_paragraph_sources(p):
    full = p.text
    ranges = [(m.start(), m.end()) for m in SRC_PAT.finditer(full)]
    if not ranges:
        return
    offset = 0
    for r in p.runs:
        rtext = r.text or ''
        if not rtext:
            offset += 0
            continue
        keep = []
        for i, ch in enumerate(rtext):
            g = offset + i
            if not any(a <= g < b for a, b in ranges):
                keep.append(ch)
        newtext = ''.join(keep)
        if newtext != rtext:
            if newtext == '':
                rp = r._element.getparent()
                if rp is not None:
                    rp.remove(r._element)
            else:
                set_text(r._element, newtext)
        offset += len(rtext)

all_paras = []
all_paras += d.paragraphs
for t in d.tables:
    for row in t.rows:
        for c in row.cells:
            all_paras += c.paragraphs
for s in d.sections:
    for part in (s.header, s.footer):
        all_paras += part.paragraphs
        for t in part.tables:
            for row in t.rows:
                for c in row.cells:
                    all_paras += c.paragraphs

for p in all_paras:
    clean_paragraph_sources(p)

# ---- gather all run elements (paragraphs, tables, header/footer) ----
run_els = []
def grab(paras):
    for p in paras:
        for r in p.runs:
            run_els.append(r._element)
grab(d.paragraphs)
for t in d.tables:
    for row in t.rows:
        for c in row.cells:
            grab(c.paragraphs)
for s in d.sections:
    for part in (s.header, s.footer):
        grab(part.paragraphs)
        for t in part.tables:
            for row in t.rows:
                for c in row.cells:
                    grab(c.paragraphs)

# drop any runs detached by source cleanup
run_els = [e for e in run_els if e.getparent() is not None]

for el in run_els:
    process_run(el)

# ---- delete answer blocks ----
allp = d.paragraphs
anchors = [p for p in allp if '参考答案与解析' in p.text]
bvol_idx = next((i for i, p in enumerate(allp) if p.text.strip() == 'B卷'), None)
ia = allp.index(anchors[0])
ib = bvol_idx
jb = allp.index(anchors[1])
remove_els = []
for p in allp[ia:ib]:          # A卷 answer header + answers + trailing blank (keep B卷 header)
    remove_els.append(p._p)
for p in allp[jb:]:             # B卷 answer header + answers to end
    remove_els.append(p._p)
for el in remove_els:
    parent = el.getparent()
    if parent is not None:
        parent.remove(el)

# ---- A3 double-column + margins ----
for sec in d.sections:
    sec.page_width = Mm(297)
    sec.page_height = Mm(420)
    sec.top_margin = Mm(18)
    sec.bottom_margin = Mm(18)
    sec.left_margin = Mm(18)
    sec.right_margin = Mm(18)
    sectPr = sec._sectPr
    cols = sectPr.find(qn('w:cols'))
    if cols is None:
        cols = etree.SubElement(sectPr, qn('w:cols'))
    cols.set(qn('w:num'), '2')
    cols.set(qn('w:space'), '680')   # ~12mm gap

# ---- fit images & tables to column width ----
for sh in d.inline_shapes:
    wmm = Emu(sh.width).mm
    if wmm > 118:
        ratio = sh.height / sh.width
        sh.width = Mm(118)
        sh.height = Mm(118 * ratio)

for t in d.tables:
    tblPr = t._tbl.tblPr
    tblW = tblPr.find(qn('w:tblW'))
    if tblW is None:
        tblW = etree.SubElement(tblPr, qn('w:tblW'))
    tblW.set(qn('w:w'), '5000')
    tblW.set(qn('w:type'), 'pct')

d.save(OUT)
print("SAVED. answer anchors:", len(anchors), " bvol_idx:", bvol_idx)
print("removed paragraph elements:", len(remove_els))
