# -*- coding: utf-8 -*-
import re
from docx import Document
from docx.oxml.ns import qn

SRC = r"C:/Users/蕾赛/Desktop/第一轮-4.5小时综合测试（A卷+B卷）最终.docx"
d = Document(SRC)
q_pat = re.compile(r'^(\d+)\.')
print("===== TABLES =====")
for ti, t in enumerate(d.tables):
    print(f"\n-- Table {ti}: {len(t.rows)} rows x {len(t.columns)} cols --")
    for ri, row in enumerate(t.rows):
        for ci, cell in enumerate(row.cells):
            for p in cell.paragraphs:
                txt = p.text.strip()
                if txt:
                    m = q_pat.match(txt)
                    tag = f"  [NUM {m.group(1)}]" if m else ""
                    print(f"   r{ri}c{ci}: {txt[:50]!r}{tag}")
                    if m:
                        for j,r in enumerate(p.runs):
                            print(f"        run{j}: {r.text!r}")

# Also: do any question paragraphs appear AFTER B卷 but the B卷 renumber should continue?
# Show last few question paras around B卷 boundary
print("\n===== AROUND B卷 (para 170-180) =====")
for i in range(170, 181):
    if i < len(d.paragraphs):
        p = d.paragraphs[i]
        m = q_pat.match(p.text)
        print(i, repr(p.text[:40]), 'NUM' if m else '')

# Count headings fonts
print("\n===== HEADING FONTS =====")
from collections import Counter
c = Counter()
for p in d.paragraphs:
    if p.style.name in ('Heading 1','Heading 2'):
        for r in p.runs:
            rp = r._element.find(qn('w:rPr'))
            ea = None
            if rp is not None:
                rf = rp.find(qn('w:rFonts'))
                if rf is not None: ea = rf.get(qn('w:eastAsia'))
            c[ea]+=1
print("heading run eastAsia fonts:", dict(c))
