# -*- coding: utf-8 -*-
"""
Final restyle of 第一轮-4.5小时综合测试（A卷+B卷）最终.docx
 1) Re-number questions continuously 1..N in document order (A卷 then B卷).
 2) Fonts: every run -> eastAsia=仿宋, ascii/hAnsi/cs=Times New Roman; strip w:hint.
 3) B卷: keep continuous (no page jump); insert two blank lines before it.
"""
import re
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SRC = r"C:/Users/蕾赛/Desktop/第一轮-4.5小时综合测试（A卷+B卷）最终（备份）.docx"
OUT = r"C:/Users/蕾赛/Desktop/第一轮-4.5小时综合测试（A卷+B卷）最终.docx"
FANG = "仿宋"
TNR = "Times New Roman"

d = Document(SRC)

# ----------------------------------------------------------------------
# 1) Fonts — global walk over ALL parts (document, headers, footers, textboxes)
# ----------------------------------------------------------------------
def set_run_fonts(r_el):
    rPr = r_el.find(qn('w:rPr'))
    if rPr is None:
        rPr = OxmlElement('w:rPr')
        r_el.insert(0, rPr)
    rf = rPr.find(qn('w:rFonts'))
    if rf is None:
        rf = OxmlElement('w:rFonts')
        rPr.insert(0, rf)
    rf.set(qn('w:eastAsia'), FANG)
    rf.set(qn('w:ascii'), TNR)
    rf.set(qn('w:hAnsi'), TNR)
    rf.set(qn('w:cs'), TNR)
    if qn('w:hint') in rf.attrib:
        del rf.attrib[qn('w:hint')]

count = 0
for part in d.part.package.iter_parts():
    root = getattr(part, 'element', None)
    if root is None:
        continue
    for r in root.iter(qn('w:r')):
        set_run_fonts(r)
        count += 1

# docDefaults so any future run inherits too
styles_el = d.styles.element
dd = styles_el.find(qn('w:docDefaults'))
if dd is None:
    dd = OxmlElement('w:docDefaults'); styles_el.insert(0, dd)
rpd = dd.find(qn('w:rPrDefault'))
if rpd is None:
    rpd = OxmlElement('w:rPrDefault'); dd.append(rpd)
rpr = rpd.find(qn('w:rPr'))
if rpr is None:
    rpr = OxmlElement('w:rPr'); rpd.append(rpr)
rfd = rpr.find(qn('w:rFonts'))
if rfd is None:
    rfd = OxmlElement('w:rFonts'); rpr.insert(0, rfd)
rfd.set(qn('w:eastAsia'), FANG)
rfd.set(qn('w:ascii'), TNR)
rfd.set(qn('w:hAnsi'), TNR)
rfd.set(qn('w:cs'), TNR)
print(f"font pass: {count} runs processed (incl. headers/footers/textboxes)")

# ----------------------------------------------------------------------
# 2) Re-number questions continuously 1..N (document order, paragraphs only)
# ----------------------------------------------------------------------
q_pat = re.compile(r'^(\d+)\.')
new_num = 0
mapping = []
for p in d.paragraphs:
    m = q_pat.match(p.text)
    if not m:
        continue
    old = int(m.group(1))
    new_num += 1
    # find k = first run whose cumulative text matches a leading "数字." pattern
    acc = ''
    k = -1
    for ri, r in enumerate(p.runs):
        acc += r.text
        if re.match(r'^\s*\d+\.', acc):
            k = ri
            break
    if k < 0:
        continue
    # length (in acc) of the leading number prefix incl. dots/space
    mfull = re.match(r'^\s*\d+\.+\s*', acc)
    numlen = mfull.end()
    pre = sum(len(p.runs[j].text) for j in range(k))   # chars before run k
    suffix = p.runs[k].text[numlen - pre:]             # text after the number within run k
    for j in range(k):
        p.runs[j].text = ''
    p.runs[k].text = f"{new_num}." + suffix
    mapping.append((old, new_num, p.text[:42]))

# ----------------------------------------------------------------------
# 3) B卷: ensure continuous (no nextPage) + two blank lines before it
# ----------------------------------------------------------------------
bv = None
for p in d.paragraphs:
    if p.text.strip() == 'B卷':
        bv = p
        break
assert bv is not None, "B卷 header not found"
bp = bv._element
parent = bp.getparent()
idx = list(parent).index(bp)
# remove any nextPage break type on the sectPr that starts B卷's section
for _ in range(2):
    parent.insert(idx, OxmlElement('w:p'))   # two blank lines
# ensure the section break before B卷 is continuous
for p in d.paragraphs:
    ppr = p._element.find(qn('w:pPr'))
    if ppr is not None:
        sp = ppr.find(qn('w:sectPr'))
        if sp is not None:
            tp = sp.find(qn('w:type'))
            if tp is not None and tp.get(qn('w:val')) == 'nextPage':
                sp.remove(tp)
print(f"renumber: {new_num} questions -> 1..{new_num}; B卷 blank lines inserted; page jump removed")

d.save(OUT)
print("SAVED. mapping sample (old->new):", mapping[:5], "...", mapping[-3:])
