# -*- coding: utf-8 -*-
import re, zipfile
from docx import Document
from docx.oxml.ns import qn

NEW = r"C:/Users/蕾赛/Desktop/第一轮-4.5小时综合测试（A卷+B卷）最终.docx"
BAK = r"C:/Users/蕾赛/Desktop/第一轮-4.5小时综合测试（A卷+B卷）最终（备份）.docx"
FANG="仿宋"; TNR="Times New Roman"

dn = Document(NEW)

# ---- 1) question number sequence ----
q_pat = re.compile(r'^(\d+)\.')
nums = []
bvol_idx = None
for i,p in enumerate(dn.paragraphs):
    if p.text.strip()=='B卷': bvol_idx=i
    m = q_pat.match(p.text)
    if m: nums.append((i,int(m.group(1))))
print("question paras:", len(nums))
seq = [n for _,n in nums]
ok = seq == list(range(1, len(seq)+1))
print("sequence 1..N contiguous:", ok)
if not ok:
    print("  ->", seq)
# show A卷 vs B卷 split
print("B卷 header at para index:", bvol_idx)
print("first B卷 question number (after B卷):", [n for i,n in nums if i>bvol_idx][:1])

# ---- 2) fonts: every w:r ----
bad_ea=[]; bad_ascii=[]; hint_left=[]
for part in dn.part.package.iter_parts():
    root=getattr(part,'element',None)
    if root is None: continue
    for r in root.iter(qn('w:r')):
        rp=r.find(qn('w:rPr'))
        if rp is None:
            bad_ea.append('no-rPr'); continue
        rf=rp.find(qn('w:rFonts'))
        if rf is None:
            bad_ea.append('no-rFonts'); continue
        if rf.get(qn('w:eastAsia'))!=FANG: bad_ea.append(rf.get(qn('w:eastAsia')))
        if rf.get(qn('w:ascii'))!=TNR: bad_ascii.append(rf.get(qn('w:ascii')))
        if rf.get(qn('w:hAnsi'))!=TNR: bad_ascii.append(rf.get(qn('w:hAnsi')))
        if rf.get(qn('w:cs'))!=TNR: bad_ascii.append(rf.get(qn('w:cs')))
        if qn('w:hint') in rf.attrib: hint_left.append(rf.get(qn('w:hint')))
print("\nfonts: bad eastAsia=",len(bad_ea)," bad ascii/hAnsi/cs=",len(bad_ascii)," hint_left=",len(hint_left))
if bad_ea[:3]: print("  eastAsia samples:", set(bad_ea)[:3])
if bad_ascii[:3]: print("  ascii samples:", set(bad_ascii)[:3])

# ---- 3) B卷 break + blank lines ----
# find paragraph element of B卷 and check 2 empty <w:p> before it
bv = dn.paragraphs[bvol_idx]
blanks=0
prev = bv._element.getprevious()
while prev is not None and prev.tag==qn('w:p'):
    # empty if no text
    txt=''.join(t.text or '' for t in prev.iter(qn('w:t')))
    if txt.strip()=='':
        blanks+=1
        prev=prev.getprevious()
    else:
        break
print("\nblank lines immediately before B卷:", blanks)
# any nextPage left?
np=0
for p in dn.paragraphs:
    ppr=p._element.find(qn('w:pPr'))
    if ppr is not None:
        sp=ppr.find(qn('w:sectPr'))
        if sp is not None:
            tp=sp.find(qn('w:type'))
            if tp is not None and tp.get(qn('w:val'))=='nextPage': np+=1
print("nextPage breaks remaining:", np)

# ---- 4) content unchanged except numbers ----
def stems(doc):
    out=[]
    for p in doc.paragraphs:
        t=p.text
        if q_pat.match(t):
            # strip leading number
            out.append(re.sub(r'^\s*\d+\.+\s*','',t))
    return out
sn=stems(dn); sb=stems(Document(BAK))
print("\nstems new/backup:", len(sn), len(sb))
print("stems identical (ignoring numbers):", sn==sb)
if sn!=sb:
    for a,b in zip(sn,sb):
        if a!=b: print("DIFF:",a[:40],"||",b[:40])
