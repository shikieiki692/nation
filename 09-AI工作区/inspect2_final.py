# -*- coding: utf-8 -*-
import re
from docx import Document
from docx.oxml.ns import qn

SRC = r"C:/Users/蕾赛/Desktop/第一轮-4.5小时综合测试（A卷+B卷）最终.docx"
d = Document(SRC)

# All top-level question numbers "N." (optional space, then any char)
q_pat = re.compile(r'^(\d+)\.')
print("===== ALL TOP-LEVEL QUESTION NUMBERS (order) =====")
nums = []
for i, p in enumerate(d.paragraphs):
    t = p.text
    m = q_pat.match(t)
    if m:
        nums.append((i, int(m.group(1)), t[:42]))
print("count:", len(nums))
prev = 0
for i, n, txt in nums:
    flag = '' if n == prev+1 else '   <-- GAP/nonsequential'
    print(f"para {i:3d}  # {n:2d}  {txt}{flag}")
    prev = n

# break type before B卷 (section 1)
print("\n===== BREAK DETAIL =====")
for i, p in enumerate(d.paragraphs):
    ppr = p._element.find(qn('w:pPr'))
    if ppr is not None:
        sp = ppr.find(qn('w:sectPr'))
        if sp is not None:
            tp = sp.find(qn('w:type'))
            val = tp.get(qn('w:val')) if tp is not None else 'NO_TYPE(continuous)'
            print(f"PARA {i} break={val} :: {p.text[:40]}")

# font sampling
print("\n===== FONT SAMPLE =====")
def fonts(run):
    rp = run._element.find(qn('w:rPr'))
    if rp is None: return (None,)*5
    rf = rp.find(qn('w:rFonts'))
    if rf is None: return (None,)*5
    return (rf.get(qn('w:eastAsia')), rf.get(qn('w:ascii')), rf.get(qn('w:hAnsi')), rf.get(qn('w:cs')), rf.get(qn('w:hint')))
cnt = {}
for p in d.paragraphs[:120]:
    for r in p.runs:
        ea, a, h, c, hint = fonts(r)
        key = (ea, a, hint)
        cnt[key] = cnt.get(key, 0) + 1
for k, v in sorted(cnt.items(), key=lambda x: -x[1])[:12]:
    print(k, '->', v)

# tables count + images
import zipfile
z = zipfile.ZipFile(SRC)
names = z.namelist()
imgs = [n for n in names if n.startswith('word/media/')]
print("\nimages:", len(imgs))
print("tables in document:", len(d.tables))
print("total paragraphs:", len(d.paragraphs))
