# -*- coding: utf-8 -*-
"""
Targeted changes on the current Desktop exam file:
1) Make B卷 force a new page  -> flip the A卷->B卷 section break to nextPage.
2) Add "第 X 页 / 共 Y 页" footer page numbers to ALL 4 sections
   (numbers in Times New Roman, Chinese in 宋体), centered.
Does NOT touch any question text, columns, or styling.
"""
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

PATH = r"C:/Users/蕾赛/Desktop/第一轮-4.5小时综合测试（A卷+B卷）.docx"
TNR = "Times New Roman"
CJK = "宋体"

d = Document(PATH)
paras = d.paragraphs

# ---------------------------------------------------------------- 1) BREAK
def find(pred, start=0):
    for i in range(start, len(paras)):
        if pred(paras[i]):
            return i
    return None

def get_sectpr(p):
    pPr = p._p.find(qn('w:pPr'))
    if pPr is None:
        return None
    return pPr.find(qn('w:sectPr'))

idx_bvol = find(lambda p: p.style.name == 'Heading 1' and p.text.strip() == 'B卷')
assert idx_bvol is not None, "B卷 header not found"

# locate the section break stored in the last A卷 question paragraph
break_sectpr = None
break_idx = None
for i in range(idx_bvol - 1, -1, -1):
    sp = get_sectpr(paras[i])
    if sp is not None:
        break_sectpr = sp
        break_idx = i
        break
assert break_sectpr is not None, "A卷->B卷 break not found"

# flip continuous -> nextPage
for child in list(break_sectpr):
    if child.tag == qn('w:type'):
        break_sectpr.remove(child)
typ = OxmlElement('w:type')
typ.set(qn('w:val'), 'nextPage')
break_sectpr.append(typ)
print("Flipped break at para", break_idx, "to nextPage")

# ---------------------------------------------------------------- 2) FOOTER
def add_field(run, instr):
    begin = OxmlElement('w:fldChar'); begin.set(qn('w:fldCharType'), 'begin')
    it = OxmlElement('w:instrText'); it.set(qn('xml:space'), 'preserve'); it.text = instr
    end = OxmlElement('w:fldChar'); end.set(qn('w:fldCharType'), 'end')
    run._r.append(begin); run._r.append(it); run._r.append(end)

def style_run(r):
    r.font.name = TNR
    r.font.size = Pt(9)
    rPr = r._element.get_or_add_rPr()
    rf = rPr.find(qn('w:rFonts'))
    if rf is None:
        rf = OxmlElement('w:rFonts'); rPr.append(rf)
    rf.set(qn('w:ascii'), TNR)
    rf.set(qn('w:hAnsi'), TNR)
    rf.set(qn('w:cs'), TNR)
    rf.set(qn('w:eastAsia'), CJK)
    # strip any eastAsia hint that could misrender symbols
    if qn('w:hint') in rf.attrib:
        del rf.attrib[qn('w:hint')]

nsec = 0
for s in d.sections:
    footer = s.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    # clear any leftover runs
    for r in list(fp.runs):
        r._r.getparent().remove(r._r)
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = fp.add_run("第 ")
    add_field(fp.add_run(), " PAGE ")
    r2 = fp.add_run(" 页 / 共 ")
    add_field(fp.add_run(), " NUMPAGES ")
    r3 = fp.add_run(" 页")
    for r in fp.runs:
        style_run(r)
    nsec += 1

print("Footers added to", nsec, "sections")
d.save(PATH)
print("SAVED OK")
