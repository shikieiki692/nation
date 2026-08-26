# -*- coding: utf-8 -*-
import re, copy
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SRC = r"C:/Obsidion/妙妙屋/00-首页/学生讲义Word/第一轮-4.5小时综合测试（A卷高考预赛+B卷竞赛）.docx"
OUT = r"C:/Obsidion/妙妙屋/00-首页/学生讲义Word/第一轮-4.5小时综合测试（A卷高考预赛+B卷竞赛）.docx"

doc = Document(SRC)

# ---------- palette ----------
BLACK = RGBColor(0x00, 0x00, 0x00)
GRAY  = RGBColor(0x66, 0x66, 0x66)
HEAD_FONT = "黑体"      # SimHei
BODY_FONT = "宋体"      # SimSun
TITLE_SIZE = 16
PART_SIZE  = 15
SECT_SIZE  = 13.5
BODY_SIZE  = 12
CAPTION_SIZE = 9
TABLE_SIZE = 10.5
SHADE_PART = "DCE6F1"   # light blue (prints faint gray)
SHADE_SECT = "EAF1FB"
SHADE_HDR  = "D9E2F3"
SHADE_TITLE= "F2F6FC"

LATIN = "Times New Roman"   # Latin / digits / symbols -> Times New Roman

def set_run_font(run, east_asia=None, size=None, bold=None, color=None, italic=None):
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts'); rPr.append(rFonts)
    if east_asia:
        rFonts.set(qn('w:eastAsia'), east_asia)
    rFonts.set(qn('w:ascii'), LATIN)
    rFonts.set(qn('w:hAnsi'), LATIN)
    rFonts.set(qn('w:cs'), LATIN)
    if size is not None: run.font.size = Pt(size)
    if bold is not None: run.font.bold = bold
    if italic is not None: run.font.italic = italic
    if color is not None: run.font.color.rgb = color

def set_para_shading(p, fill):
    pPr = p._p.get_or_add_pPr()
    shd = pPr.find(qn('w:shd'))
    if shd is None:
        shd = OxmlElement('w:shd'); pPr.append(shd)
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), fill)

def set_para_bottom_border(p, color="4472C4", sz="8"):
    pPr = p._p.get_or_add_pPr()
    pBdr = pPr.find(qn('w:pBdr'))
    if pBdr is None:
        pBdr = OxmlElement('w:pBdr'); pPr.append(pBdr)
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single'); bottom.set(qn('w:sz'), sz)
    bottom.set(qn('w:space'), '4'); bottom.set(qn('w:color'), color)
    pBdr.append(bottom)

def set_align(p, align):
    p.alignment = align

def set_spacing(p, before=0, after=4, line=1.5, rule=WD_LINE_SPACING.MULTIPLE):
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing_rule = rule
    pf.line_spacing = line

# ---------- page setup ----------
sec = doc.sections[0]
sec.page_width = Cm(21.0); sec.page_height = Cm(29.7)
sec.top_margin = Cm(2.0); sec.bottom_margin = Cm(2.0)
sec.left_margin = Cm(2.2); sec.right_margin = Cm(2.0)

# default Normal font
normal = doc.styles['Normal']
normal.font.name = LATIN
normal.font.size = Pt(BODY_SIZE)
_npr = normal.element.get_or_add_rPr()
_nrf = _npr.find(qn('w:rFonts'))
if _nrf is None:
    _nrf = OxmlElement('w:rFonts'); _npr.append(_nrf)
_nrf.set(qn('w:ascii'), LATIN)
_nrf.set(qn('w:hAnsi'), LATIN)
_nrf.set(qn('w:cs'), LATIN)
_nrf.set(qn('w:eastAsia'), BODY_FONT)

# ---------- helpers for runs ----------
def all_runs(p):
    return list(p.runs)

QRE = re.compile(r'^(\d+\.）（[^）]*）)')     # 1.（2024 山东卷·题4，★★）
QRE2 = re.compile(r'^(\d+\.）)')            # fallback 1.）

def style_question_item(p):
    set_align(p, WD_ALIGN_PARAGRAPH.LEFT)
    set_spacing(p, before=2, after=4, line=1.5)
    txt = p.text
    for r in p.runs:
        set_run_font(r, east_asia=BODY_FONT, size=BODY_SIZE, bold=False, color=BLACK)
    m = QRE.match(txt) or QRE2.match(txt)
    if m:
        target = m.group(1)
        # bold leading group
        done = 0
        for r in p.runs:
            if done >= len(target): break
            rem = len(target) - done
            if len(r.text) <= rem:
                set_run_font(r, east_asia=HEAD_FONT, size=BODY_SIZE, bold=True, color=BLACK)
                done += len(r.text)
            else:
                # split run: keep head bold, move tail into a new following run
                head = r.text[:rem]; tail = r.text[rem:]
                r.text = head
                set_run_font(r, east_asia=HEAD_FONT, size=BODY_SIZE, bold=True, color=BLACK)
                from docx.text.run import Run
                new_r = OxmlElement('w:r')
                r._element.addnext(new_r)
                new_run = Run(new_r, r._parent)
                new_run.text = tail
                set_run_font(new_run, east_asia=BODY_FONT, size=BODY_SIZE, bold=False, color=BLACK)
                done += rem
    set_para_shading(p, "FFFFFF")

def style_body(p, caption=False):
    set_align(p, WD_ALIGN_PARAGRAPH.LEFT)
    if caption:
        set_spacing(p, before=1, after=6, line=1.3)
        for r in p.runs:
            set_run_font(r, east_asia=BODY_FONT, size=CAPTION_SIZE, bold=False, color=GRAY, italic=True)
    else:
        set_spacing(p, before=1, after=4, line=1.5)
        for r in p.runs:
            set_run_font(r, east_asia=BODY_FONT, size=BODY_SIZE, bold=False, color=BLACK)

# ---------- iterate paragraphs ----------
paragraphs = doc.paragraphs
title_done = False
for idx, p in enumerate(paragraphs):
    txt = p.text.strip()
    style = p.style.name
    if not txt:
        # keep empty paragraphs but ensure spacing minimal
        continue

    is_heading1 = (style == 'Heading 1')
    is_heading2 = (style == 'Heading 2')

    if is_heading1 and (idx == 0 or ('综合测试' in txt and '卷' in txt and p.alignment == WD_ALIGN_PARAGRAPH.CENTER)):
        # TITLE
        set_align(p, WD_ALIGN_PARAGRAPH.CENTER)
        set_spacing(p, before=4, after=8, line=1.2)
        for r in p.runs:
            set_run_font(r, east_asia=HEAD_FONT, size=TITLE_SIZE, bold=True, color=BLACK)
        set_para_shading(p, SHADE_TITLE)
        set_para_bottom_border(p, color="2E5496", sz="12")
        title_done = True
        continue

    if is_heading1 and ('卷' in txt):
        # A卷 / B卷 major part header
        set_align(p, WD_ALIGN_PARAGRAPH.CENTER)
        set_spacing(p, before=10, after=8, line=1.2)
        for r in p.runs:
            set_run_font(r, east_asia=HEAD_FONT, size=PART_SIZE, bold=True, color=BLACK)
        set_para_shading(p, SHADE_PART)
        set_para_bottom_border(p, color="2E5496", sz="10")
        continue

    if is_heading1:
        # any other Heading 1 -> treat as major part
        set_align(p, WD_ALIGN_PARAGRAPH.CENTER)
        set_spacing(p, before=10, after=8, line=1.2)
        for r in p.runs:
            set_run_font(r, east_asia=HEAD_FONT, size=PART_SIZE, bold=True, color=BLACK)
        set_para_shading(p, SHADE_PART)
        set_para_bottom_border(p, color="2E5496", sz="10")
        continue

    if is_heading2:
        # 大题 section header: 一、二、三……
        set_align(p, WD_ALIGN_PARAGRAPH.LEFT)
        set_spacing(p, before=10, after=6, line=1.2)
        for r in p.runs:
            set_run_font(r, east_asia=HEAD_FONT, size=SECT_SIZE, bold=True, color=BLACK)
        set_para_shading(p, SHADE_SECT)
        set_para_bottom_border(p, color="8EAADB", sz="6")
        continue

    # caption
    if txt.startswith('图：'):
        style_body(p, caption=True)
        continue

    # question item
    if QRE.match(txt) or QRE2.match(txt):
        style_question_item(p)
        continue

    # default body
    style_body(p, caption=False)

# ---------- tables ----------
def shade_cell(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn('w:shd'))
    if shd is None:
        shd = OxmlElement('w:shd'); tcPr.append(shd)
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), fill)

def set_cell_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement('w:tblBorders')
    for edge in ('top','left','bottom','right','insideH','insideV'):
        e = OxmlElement(f'w:{edge}')
        e.set(qn('w:val'), 'single'); e.set(qn('w:sz'), '4')
        e.set(qn('w:space'), '0'); e.set(qn('w:color'), '808080')
        borders.append(e)
    tblPr.append(borders)

for t in doc.tables:
    set_cell_borders(t)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    # header row shading + bold
    for j, cell in enumerate(t.rows[0].cells):
        shade_cell(cell, SHADE_HDR)
        for p in cell.paragraphs:
            for r in p.runs:
                set_run_font(r, east_asia=HEAD_FONT, size=TABLE_SIZE, bold=True, color=BLACK)
    # body cells
    for ri in range(1, len(t.rows)):
        for cell in t.rows[ri].cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    set_run_font(r, east_asia=BODY_FONT, size=TABLE_SIZE, bold=False, color=BLACK)

# ---------- header + footer ----------
def add_field(paragraph, field):
    run = paragraph.add_run()
    fld1 = OxmlElement('w:fldChar'); fld1.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText'); instr.set(qn('xml:space'), 'preserve'); instr.text = f' {field} '
    fld2 = OxmlElement('w:fldChar'); fld2.set(qn('w:fldCharType'), 'end')
    run._r.append(fld1); run._r.append(instr); run._r.append(fld2)
    return run

# header: slim centered title
hdr = sec.header
hp = hdr.paragraphs[0]
hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
hr = hp.add_run("第一轮 · 4.5h 综合测试（A卷 高考预赛 + B卷 竞赛）")
set_run_font(hr, east_asia=BODY_FONT, size=9, bold=False, color=GRAY)
set_para_bottom_border(hp, color="BFBFBF", sz="4")

# footer: page numbers
ftr = sec.footer
fp = ftr.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r1 = fp.add_run("第 "); set_run_font(r1, east_asia=BODY_FONT, size=9, color=GRAY)
add_field(fp, "PAGE")
r2 = fp.add_run(" 页 / 共 "); set_run_font(r2, east_asia=BODY_FONT, size=9, color=GRAY)
add_field(fp, "NUMPAGES")
r3 = fp.add_run(" 页"); set_run_font(r3, east_asia=BODY_FONT, size=9, color=GRAY)

doc.save(OUT)
print("SAVED:", OUT)
print("paragraphs:", len(paragraphs), "tables:", len(doc.tables))
