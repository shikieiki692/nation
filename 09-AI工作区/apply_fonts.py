# -*- coding: utf-8 -*-
"""Surgical font pass on the print version:
   Keep each run's Chinese font (eastAsia) but set Latin/digits to Times New Roman.
   Does NOT touch any text, vertAlign (subscript/superscript), or structure.
"""
import copy
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SRC = r"C:/Obsidion/妙妙屋/00-首页/学生讲义Word/第一轮-4.5小时综合测试（A卷高考预赛+B卷竞赛）_打印版.docx"
LATIN = "Times New Roman"
DEFAULT_EAST = "宋体"

doc = Document(SRC)

def get_eastasia(run):
    rPr = run._element.find(qn('w:rPr'))
    if rPr is None: return None
    rf = rPr.find(qn('w:rFonts'))
    if rf is None: return None
    return rf.get(qn('w:eastAsia'))

def set_split(run, east_asia):
    rPr = run._element.get_or_add_rPr()
    rf = rPr.find(qn('w:rFonts'))
    if rf is None:
        rf = OxmlElement('w:rFonts'); rPr.append(rf)
    rf.set(qn('w:ascii'), LATIN)
    rf.set(qn('w:hAnsi'), LATIN)
    rf.set(qn('w:cs'), LATIN)
    rf.set(qn('w:eastAsia'), east_asia)

def process_para(p):
    for r in p.runs:
        ea = get_eastasia(r) or DEFAULT_EAST
        set_split(r, ea)

# 1) body paragraphs
for p in doc.paragraphs:
    process_para(p)

# 2) tables
for t in doc.tables:
    for row in t.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                process_para(p)

# 3) header & footer
sec = doc.sections[0]
for hp in sec.header.paragraphs:
    process_para(hp)
for fp in sec.footer.paragraphs:
    process_para(fp)

# 4) text boxes (drawing txbxContent) if any
def walk_txbx(el):
    from docx.oxml.ns import qn as _qn
    found = 0
    for txbx in el.iter(_qn('w:txbxContent')):
        for p in txbx.iter(_qn('w:p')):
            # build a lightweight run iterator
            for r in p.iter(_qn('w:r')):
                from docx.text.run import Run
                run = Run(r, None)
                ea = get_eastasia(run) or DEFAULT_EAST
                set_split(run, ea)
                found += 1
    return found

tb = walk_txbx(doc.element.body)
tb2 = 0
for s in doc.sections:
    try: tb2 += walk_txbx(s.header._element)
    except Exception: pass
    try: tb2 += walk_txbx(s.footer._element)
    except Exception: pass

# 5) Normal style default fonts
normal = doc.styles['Normal']
_npr = normal.element.get_or_add_rPr()
_nrf = _npr.find(qn('w:rFonts'))
if _nrf is None:
    _nrf = OxmlElement('w:rFonts'); _npr.append(_nrf)
_nrf.set(qn('w:ascii'), LATIN)
_nrf.set(qn('w:hAnsi'), LATIN)
_nrf.set(qn('w:cs'), LATIN)
_nrf.set(qn('w:eastAsia'), DEFAULT_EAST)

doc.save(SRC)
print("FONT PASS DONE")
print("textboxes processed runs:", tb + tb2)
