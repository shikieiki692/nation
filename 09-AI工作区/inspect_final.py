# -*- coding: utf-8 -*-
import re
from docx import Document
from docx.oxml.ns import qn

SRC = r"C:/Users/蕾赛/Desktop/第一轮-4.5小时综合测试（A卷+B卷）最终.docx"
d = Document(SRC)

def font_of(run):
    rp = run._element.find(qn('w:rPr'))
    if rp is None:
        return (None, None, None, None)
    rf = rp.find(qn('w:rFonts'))
    if rf is None:
        return (None, None, None, None)
    ea = rf.get(qn('w:eastAsia'))
    asc = rf.get(qn('w:ascii'))
    hA = rf.get(qn('w:hAnsi'))
    cs = rf.get(qn('w:cs'))
    hint = rf.get(qn('w:hint'))
    return (ea, asc, hA, cs, hint)

# 1) dump first ~15 and the question-number paragraphs
print("===== HEAD (first 14 paras) =====")
for i, p in enumerate(d.paragraphs[:14]):
    print(i, repr(p.text[:70]), '| style=', p.style.name)

# 2) find question-number paragraphs like '1.' '2.' at start
print("\n===== QUESTION-NUMBERED PARAS (top-level 数字.) =====")
q_pat = re.compile(r'^\s*(\d+)\.\s')
hdr_pat = re.compile(r'^[一-龥]?卷$|^一、|^二、|^三、|^四、|^五、|^六、')
questions = []
for i, p in enumerate(d.paragraphs):
    t = p.text.strip()
    m = q_pat.match(t)
    if m:
        questions.append((i, int(m.group(1)), t[:50]))
print("total numbered paragraphs:", len(questions))
for i, num, txt in questions[:60]:
    print(i, num, txt)

# 3) list A卷 / B卷 / section headers
print("\n===== VOLUME / SECTION HEADERS =====")
for i, p in enumerate(d.paragraphs):
    t = p.text.strip()
    if t in ('A卷', 'B卷') or re.match(r'^[一二三四五六七]、', t):
        print(i, repr(t), '| style=', p.style.name)

# 4) section breaks
print("\n===== SECTIONS / BREAKS =====")
for i, s in enumerate(d.sections):
    sp = s._sectPr
    # break type before this section (look at preceding paragraph's pPr/w:rPr? actually break is in a paragraph's pPr)
    cols = sp.find(qn('w:cols'))
    ncol = cols.get(qn('w:num')) if cols is not None else '?'
    pg = sp.find(qn('w:pgSz'))
    print(f"Section {i}: cols={ncol}")

# find paragraphs containing sectPr (section breaks)
for i, p in enumerate(d.paragraphs):
    ppr = p._element.find(qn('w:pPr'))
    if ppr is not None:
        sp = ppr.find(qn('w:sectPr'))
        if sp is not None:
            typ = None
            brk = sp.find(qn('w:type'))
            # type of break is in w:sectPr directly? Actually break type is stored as child <w:type val="nextPage"/> inside sectPr when it's a section break paragraph
            tp = sp.find(qn('w:type'))
            tval = tp.get(qn('w:val')) if tp is not None else 'continuous?'
            print(f"  PARA {i} has sectPr with break type: {tval} :: {p.text[:30]}")
