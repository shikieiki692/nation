# -*- coding: utf-8 -*-
"""
Restructure the A3 double-column exam into a clean printed-paper layout:
- Title and A卷/B卷 volume headers span the FULL A3 width (single-column sections)
- Questions flow in two columns (A3, ~12mm gap)
- Breathing room between questions and between option lines
- Prettier headings: shaded volume/section headers + bottom rules
Operates on the already-cleaned file (stars/sources/answers removed, TNR fonts).
"""
import re, copy
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt, Mm
from docx.oxml import OxmlElement
from lxml import etree

PATH = r"C:/Users/蕾赛/Desktop/第一轮-4.5小时综合测试（A卷+B卷）.docx"
TNR = "Times New Roman"
XMLSPACE = "{http://www.w3.org/XML/1998/namespace}space"

# A3 portrait — Word stores pgSz/pgMar/cols-space in TWIPS (1mm = 1440/25.4 twips)
def mm_tw(mm):
    return int(round(mm * 1440.0 / 25.4))
A3_W = mm_tw(297); A3_H = mm_tw(420)
M = mm_tw(18)              # margins ~1021 twips
COL_GAP = mm_tw(12)        # ~680 twips (~12mm)

d = Document(PATH)

# ----------------------------------------------------------------------------
# 1. Identify key paragraphs by text/style
# ----------------------------------------------------------------------------
paras = d.paragraphs
def find(pred, start=0):
    for i in range(start, len(paras)):
        if pred(paras[i]):
            return i
    return None

idx_title = find(lambda p: p.style.name == 'Heading 1' and '综合测试' in p.text)
idx_avol  = find(lambda p: p.style.name == 'Heading 1' and p.text.strip() == 'A卷', idx_title)
idx_bvol  = find(lambda p: p.style.name == 'Heading 1' and p.text.strip() == 'B卷', idx_avol)
# last A卷 question = paragraph just before B卷 volume header
idx_lastA = idx_bvol - 1

print("idx_title", idx_title, "idx_avol", idx_avol, "idx_lastA", idx_lastA, "idx_bvol", idx_bvol)

# ----------------------------------------------------------------------------
# 2. Rebuild sections: 4 sections
#    S1 single(col1): title + A卷 header  (break after idx_avol)
#    S2 two-col      : A卷 questions        (break after idx_lastA)
#    S3 single       : B卷 header           (break after idx_bvol)
#    S4 two-col      : B卷 questions        (body final sectPr)
# ----------------------------------------------------------------------------
def build_sectPr(num, continuous):
    sp = OxmlElement('w:sectPr')
    if continuous:
        t = OxmlElement('w:type'); t.set(qn('w:val'), 'continuous'); sp.append(t)
    pg = OxmlElement('w:pgSz')
    pg.set(qn('w:w'), str(A3_W)); pg.set(qn('w:h'), str(A3_H)); pg.set(qn('w:orient'), 'portrait')
    sp.append(pg)
    mg = OxmlElement('w:pgMar')
    mg.set(qn('w:top'), str(M)); mg.set(qn('w:bottom'), str(M))
    mg.set(qn('w:left'), str(M)); mg.set(qn('w:right'), str(M))
    mg.set(qn('w:header'), '720'); mg.set(qn('w:footer'), '720'); mg.set(qn('w:gutter'), '0')
    sp.append(mg)
    cols = OxmlElement('w:cols')
    cols.set(qn('w:num'), str(num))
    if num > 1:
        cols.set(qn('w:space'), str(COL_GAP))
    sp.append(cols)
    return sp

# remove any existing mid-document sectPr
for p in paras:
    pPr = p._p.find(qn('w:pPr'))
    if pPr is not None:
        s = pPr.find(qn('w:sectPr'))
        if s is not None:
            pPr.remove(s)

def set_para_break(p, sectPr):
    pPr = p._p.get_or_add_pPr()
    old = pPr.find(qn('w:sectPr'))
    if old is not None:
        pPr.remove(old)
    pPr.append(sectPr)   # must be last child of pPr

set_para_break(paras[idx_avol],  build_sectPr(1, True))   # end S1
set_para_break(paras[idx_lastA], build_sectPr(2, True))   # end S2
set_para_break(paras[idx_bvol],  build_sectPr(1, True))   # end S3

# body final sectPr -> S4 (two columns, no type)
body = d.element.body
old_final = body.find(qn('w:sectPr'))
if old_final is not None:
    body.remove(old_final)
final = build_sectPr(2, False)
body.append(final)

print("sections now:", len(d.sections))

# ----------------------------------------------------------------------------
# 3. Spacing + heading sizes + shading/borders
# ----------------------------------------------------------------------------
def set_shade(p, fill):
    pPr = p._p.get_or_add_pPr()
    shd = pPr.find(qn('w:shd'))
    if shd is None:
        shd = OxmlElement('w:shd'); pPr.append(shd)
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), fill)

def set_bottom_border(p, sz, color):
    pPr = p._p.get_or_add_pPr()
    pBdr = pPr.find(qn('w:pBdr'))
    if pBdr is None:
        pBdr = OxmlElement('w:pBdr'); pPr.append(pBdr)
    b = pBdr.find(qn('w:bottom'))
    if b is None:
        b = OxmlElement('w:bottom'); pBdr.append(b)
    b.set(qn('w:val'), 'single'); b.set(qn('w:sz'), str(sz)); b.set(qn('w:space'), '2'); b.set(qn('w:color'), color)

def set_run_size_bold(p, size, bold=True):
    for r in p.runs:
        r.font.size = Pt(size)
        r.font.bold = bold

QSTART = re.compile(r'^\d+[\.．]')
OPT2 = ('A.', 'B.', 'C.', 'D.', 'E.')

# global comfortable line spacing on the common styles
for st in ('Normal', 'Body Text', 'First Paragraph', 'Heading 1', 'Heading 2'):
    try:
        s = d.styles[st]
        s.paragraph_format.line_spacing = 1.4
    except KeyError:
        pass

title_done = False
for p in paras:
    txt = p.text
    st = p.style.name
    if st == 'Heading 1':
        if not title_done and '综合测试' in txt:
            # main title
            title_done = True
            p.alignment = 1  # center
            set_run_size_bold(p, 20)
            set_bottom_border(p, 18, '2E75B6')
            p.paragraph_format.space_after = Pt(10)
            p.paragraph_format.space_before = Pt(2)
        else:
            # volume header (A卷 / B卷)
            p.alignment = 1
            set_run_size_bold(p, 15)
            set_shade(p, 'DEEAF6')
            set_bottom_border(p, 12, '2E75B6')
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(8)
    elif st == 'Heading 2':
        set_run_size_bold(p, 13)
        set_shade(p, 'F2F2F2')
        set_bottom_border(p, 6, 'BFBFBF')
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(6)
    else:
        t = txt.strip()
        if QSTART.match(t):
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(2)
        elif t[:2] in OPT2 or (t and t[0] in 'ABCD' and t[1:2] in '．.'):
            p.paragraph_format.space_after = Pt(3)

# ----------------------------------------------------------------------------
# 4. Make sure the 3 tables & images still fit the narrower column nicely
# ----------------------------------------------------------------------------
for sh in d.inline_shapes:
    wmm = sh.width.mm if hasattr(sh.width, 'mm') else None
    if wmm is None:
        continue
    if wmm > 115:
        ratio = sh.height / sh.width
        sh.width = Mm(115)
        sh.height = Mm(115 * ratio)

for t in d.tables:
    tblPr = t._tbl.tblPr
    tblW = tblPr.find(qn('w:tblW'))
    if tblW is None:
        tblW = etree.SubElement(tblPr, qn('w:tblW'))
    tblW.set(qn('w:w'), '5000'); tblW.set(qn('w:type'), 'pct')

d.save(PATH)
print("SAVED OK")
